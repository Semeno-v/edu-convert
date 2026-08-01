"""Чтение старых документов Word (python-docx).

Три задачи (ТЗ §4 Этап 2.1/2.3/2.5):

1. :meth:`WordExtractor.extract_index` — определить индекс дисциплины
   (титул/первая таблица).
2. :meth:`WordExtractor.extract_old_numbers` — вытащить **старые** числа из
   первой таблицы (по ключевым словам шапки) — только для отчёта о расхождениях.
3. :meth:`WordExtractor.extract` — конечный автомат по абзацам: при встрече
   заголовка включается «режим записи» блока до следующего заголовка; контент
   копируется в сериализуемый IR (:class:`ContentBlocks`).

Заголовки распознаются **по тексту** (нумерация + заглавная буква), а не по
стилям — стили в исходных документах непоследовательны (ТЗ §4 Этап 2.5).
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocument
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from app.core.exceptions import DocumentParseError
from app.core.models import (
    ContentBlock,
    ContentBlocks,
    ContentElement,
    DocType,
    ElementKind,
    OldNumbers,
    RichParagraph,
    RichRun,
    RichTable,
    RichTableCell,
    RichTableRow,
)
from app.core.normalizer import (
    as_int,
    index_from_text,
    normalize_text,
    parse_hours_ze,
)

# --------------------------------------------------------------------------- #
#  Правила сопоставления заголовков с каноническими ключами блоков.
#  ⚙️ Это главная предметная настройка: порядок важен (первое совпадение).
#  Можно расширять под особенности конкретных РПД/ФОС, не трогая логику обхода.
# --------------------------------------------------------------------------- #
IGNORED_KEY = "_ignore"  # сток: заголовок закрывает блок, контент отбрасывается
BLOCK_RULES: list[tuple[str, tuple[str, ...]]] = [
    # --- РПД ---
    ("goals", ("цели освоения", "цель дисциплины")),
    ("competencies", ("компетенц", "планируемые результаты")),
    ("thematic_plan", (
        "тематическ", "траектори", "структура и содержание", "содержание дисциплины",
    )),
    ("literature_main", ("основная литература",)),
    ("literature_extra", ("дополнительная литература",)),
    ("periodicals", ("периодическ",)),
    ("methodical", ("методические указания",)),
    # «Дополнительные средства обучения (в т.ч. on-line курсы)» — тоже
    # интернет-ресурсы: их таблицы (например, «Совнет») сохраняются в §4.3.
    ("internet_resources", (
        "интернет", "электронные образовательные", "профессиональные базы",
        "справочные систем", "дополнительные средства обучения",
    )),
    ("software", ("программное обеспечение", "лицензионное")),
    ("facilities", ("материально-техническ", "технические средства", "помещения")),
    ("assessment", ("оценочные средства", "оценка качества", "система оцениван", "шкала оцениван")),
    # --- ФОС ---
    ("self_study_questions", ("самостоятельной подготовки",)),
    ("current_control", ("текущего контроля", "текущий контроль")),
    ("interim_attestation", ("промежуточной аттестации", "промежуточн")),
    ("gia", ("государственной итоговой", "итоговой аттестации", "ги(и)а")),
    # --- Стоп-заголовки: закрывают текущий блок, их контент никуда не идёт.
    # «4. Учебно-методическое обеспечение» и «4.1 Рекомендуемая литература» —
    # обёртки над 4.1.1/4.1.2 (те ловятся своими правилами); «Дополнительные
    # средства обучения» — заглушка перед «Профессиональными базами данных».
    (IGNORED_KEY, ("учебно-методическое", "рекомендуемая литература")),
]

# Блоки, существующие в ФОС: только они (и стоп-заголовки) переключают автомат
# при doc_type=FOS; заголовки с РПД-ключами (шкалы/критерии оценивания,
# «…компетенций») в ФОС закрывают блок — их место в §8 РПД, одобренные
# эталоны ФОС 2026 шкал и таблиц оценивания не содержат.
_FOS_KEYS = frozenset({"self_study_questions", "current_control", "interim_attestation", "gia"})

# Ненумерованные подзаголовки-маркеры в ФОС: тип вопросов определяет целевой
# раздел (так разложены одобренные эталоны: задачи → §1, тесты → §2).
_FOS_MARKERS: tuple[tuple[str | None, tuple[str, ...]], ...] = (
    ("current_control", (
        "развернутым ответом", "развернутый ответ", "перечень задач", "открытого типа",
    )),
    ("interim_attestation", (
        "тестовых вопрос", "тестовые вопрос", "тестовых задани", "тестовые задани",
        "закрытого типа",
    )),
    (None, ("описание шкал",)),  # None → сток (контент отбрасывается)
)

# Строка оглавления: текст + таб + номер страницы («…компетенций\t8»).
_TOC_LINE_RE = re.compile(r"\t\d{1,3}\s*$")

# Пробел после номера необязателен: распознаём и «3. Тема», и «3.Тема».
_HEADING_RE = re.compile(r"^\s*\d+(?:\.\d+)*\.?\s*\S")
_MAX_HEADING_LEN = 200


def _heading_key(text: str) -> str | None:
    """Возвращает канонический ключ блока по тексту заголовка (или None)."""
    norm = normalize_text(text)
    for key, keywords in BLOCK_RULES:
        if any(normalize_text(kw) in norm for kw in keywords):
            return key
    return None


def is_heading(text: str) -> bool:
    """Эвристика: нумерованный, относительно короткий заголовок раздела."""
    stripped = text.strip()
    if not stripped or len(stripped) > _MAX_HEADING_LEN:
        return False
    return bool(_HEADING_RE.match(stripped))


def _is_bold_numbered_heading(paragraph: Paragraph) -> bool:
    """Заголовок-список: прямой numPr + целиком жирный текст.

    Часть исходников нумерует разделы обычным Word-списком со стилем
    «List Paragraph» («Оценка качества реализации дисциплины» и т.п.) —
    цепочка стилей тут не помогает. Контентные нумерованные списки
    (вопросы, перечни) целиком жирными не бывают.
    """
    if _num_props(paragraph) is None:
        return False
    runs = [r for r in paragraph.runs if r.text.strip()]
    return bool(runs) and all(r.bold for r in runs)


def _has_heading_style(paragraph: Paragraph) -> bool:
    """Заголовок без видимого номера: нумерация/уровень структуры — в стиле.

    В части исходников номера разделов даёт Word-автонумерация стилей
    («Основная литература» со стилем на базе Heading 2) — текст абзаца числа
    не содержит, и текстовая эвристика такие заголовки не видит. Прямой
    ``numPr`` самого абзаца (нумерованные списки контента) сюда не входит.
    """
    style = paragraph.style
    seen: set[str] = set()
    while style is not None and style.style_id not in seen:
        seen.add(style.style_id)
        name = style.name or ""
        if name.startswith(("Heading", "Заголовок")):
            return True
        pPr = style.element.find(qn("w:pPr"))
        if pPr is not None and (
            pPr.find(qn("w:numPr")) is not None or pPr.find(qn("w:outlineLvl")) is not None
        ):
            return True
        style = style.base_style
    return False


# --------------------------------------------------------------------------- #
#  Преобразование python-docx → IR
# --------------------------------------------------------------------------- #


def _iter_block_items(parent: DocxDocument | _Cell):
    """Итерирует абзацы и таблицы в порядке следования в документе/ячейке."""
    parent_elm = parent.element.body if isinstance(parent, DocxDocument) else parent._tc
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def _num_props(paragraph: Paragraph) -> tuple[int, int] | None:
    """(numId, ilvl) автонумерации абзаца или None (numId=0 — нумерация снята)."""
    pPr = paragraph._p.pPr
    if pPr is None:
        return None
    numPr = pPr.find(qn("w:numPr"))
    if numPr is None:
        return None
    num_el = numPr.find(qn("w:numId"))
    nid = as_int(num_el.get(qn("w:val"))) if num_el is not None else None
    if not nid:
        return None
    ilvl_el = numPr.find(qn("w:ilvl"))
    lvl = as_int(ilvl_el.get(qn("w:val"))) if ilvl_el is not None else 0
    return nid, lvl or 0


class _NumberingModel:
    """Best-effort синтез видимых номеров автонумерации Word.

    Автонумерация живёт в ``numbering.xml`` и при переносе абзацев в другой
    документ теряется (вопросы «1.–10.» становились безномерными). Модель
    читает определения уровней (формат, шаблон ``lvlText``, старт) и ведёт
    счётчики по (numId, ilvl), отдавая текстовый префикс «1. », «a) », «– ».
    """

    _DEFAULT = ("decimal", "%1.", 1)

    def __init__(self, doc: DocxDocument) -> None:
        self._levels: dict[tuple[int, int], tuple[str, str, int]] = {}
        self._counters: dict[tuple[int, int], int] = {}
        try:
            # part_related_by объявлен возвращающим базовый Part, а .element есть
            # только у XmlPart; отсутствие атрибута тут же ловится ниже.
            root = doc.part.part_related_by(RT.NUMBERING).element  # type: ignore[attr-defined]
        except (KeyError, AttributeError):
            return
        abstract: dict[int, dict[int, tuple[str, str, int]]] = {}
        for an in root.findall(qn("w:abstractNum")):
            levels: dict[int, tuple[str, str, int]] = {}
            for lvl in an.findall(qn("w:lvl")):
                fmt_el, txt_el, start_el = (
                    lvl.find(qn(t)) for t in ("w:numFmt", "w:lvlText", "w:start")
                )
                levels[int(lvl.get(qn("w:ilvl")))] = (
                    fmt_el.get(qn("w:val")) if fmt_el is not None else "decimal",
                    txt_el.get(qn("w:val")) if txt_el is not None else "%1.",
                    as_int(start_el.get(qn("w:val"))) or 1 if start_el is not None else 1,
                )
            abstract[int(an.get(qn("w:abstractNumId")))] = levels
        for num in root.findall(qn("w:num")):
            aref = num.find(qn("w:abstractNumId"))
            if aref is None:
                continue
            nid = int(num.get(qn("w:numId")))
            for ilvl, spec in abstract.get(int(aref.get(qn("w:val"))), {}).items():
                self._levels[(nid, ilvl)] = spec

    def _render_level(self, num_id: int, level: int) -> str:
        fmt, _, start = self._levels.get((num_id, level), self._DEFAULT)
        n = self._counters.get((num_id, level), start)
        if fmt == "lowerLetter" and 1 <= n <= 26:
            return chr(ord("a") + n - 1)
        if fmt == "upperLetter" and 1 <= n <= 26:
            return chr(ord("A") + n - 1)
        return str(n)

    def is_decimal_dot(self, num_id: int, ilvl: int) -> bool:
        """Десятичная нумерация вида «N.» (кандидат на формат «Задача № N: »)."""
        fmt, lvl_text, _ = self._levels.get((num_id, ilvl), self._DEFAULT)
        return fmt != "bullet" and lvl_text.endswith(".")

    def prefix(self, num_id: int, ilvl: int, *, task_style: bool = False) -> str:
        fmt, lvl_text, start = self._levels.get((num_id, ilvl), self._DEFAULT)
        if fmt == "bullet":
            return "– "
        self._counters[(num_id, ilvl)] = self._counters.get((num_id, ilvl), start - 1) + 1
        for key in [k for k in self._counters if k[0] == num_id and k[1] > ilvl]:
            del self._counters[key]  # более глубокие уровни начинаются заново
        if task_style:
            # Формат одобренных эталонов ФОС: вопросы — «Задача № N: …».
            return f"Задача № {self._render_level(num_id, ilvl)}: "
        rendered = re.sub(
            r"%(\d)", lambda m: self._render_level(num_id, int(m.group(1)) - 1), lvl_text
        )
        if not rendered:
            return ""
        # Варианты ответов «N)» в эталонах отделены табом, прочее — пробелом.
        return f"{rendered}\t" if rendered.endswith(")") else f"{rendered} "


def _list_level(paragraph: Paragraph) -> int | None:
    """Уровень списка из нумерации OOXML (или None)."""
    pPr = paragraph._p.pPr
    if pPr is None:
        return None
    numPr = pPr.find(qn("w:numPr"))
    if numPr is None:
        return None
    ilvl = numPr.find(qn("w:ilvl"))
    if ilvl is not None:
        val = ilvl.get(qn("w:val"))
        return as_int(val) if val is not None else 0
    return 0


# Контейнеры, внутри которых лежат раны, пропускаемые paragraph.runs:
# гиперссылки, smartTag (могут быть вложенными), простые поля.
_RUN_CONTAINER_TAGS = ("w:hyperlink", "w:smartTag", "w:fldSimple", "w:ins")


def _iter_runs(paragraph: Paragraph):
    """Раны абзаца в порядке документа, включая вложенные в w:hyperlink/w:smartTag.

    ``paragraph.runs`` python-docx отдаёт только прямые ``w:r`` — текст ссылок
    («URL: …») и smartTag-значений («2 см») иначе теряется при переносе.
    """
    container_tags = {qn(tag) for tag in _RUN_CONTAINER_TAGS}

    def walk(element):
        for child in element:
            if child.tag == qn("w:r"):
                yield Run(child, paragraph)
            elif child.tag in container_tags:
                yield from walk(child)

    yield from walk(paragraph._p)


_ALIGNMENT_NAMES = {1: "center", 2: "right", 3: "justify", 0: "left"}


def para_to_rich(paragraph: Paragraph) -> RichParagraph:
    """Преобразует абзац python-docx в :class:`RichParagraph` (с форматированием)."""
    runs = [
        RichRun(
            text=run.text,
            bold=bool(run.bold),
            italic=bool(run.italic),
            underline=bool(run.underline),
        )
        for run in _iter_runs(paragraph)
        if run.text
    ]
    if not runs and paragraph.text.strip():
        runs = [RichRun(text=paragraph.text)]
    style = paragraph.style.name if paragraph.style is not None else None
    pf = paragraph.paragraph_format
    return RichParagraph(
        runs=runs,
        style=style,
        list_level=_list_level(paragraph),
        alignment=_ALIGNMENT_NAMES.get(int(pf.alignment)) if pf.alignment is not None else None,
        first_line_indent_pt=pf.first_line_indent.pt if pf.first_line_indent is not None else None,
        left_indent_pt=pf.left_indent.pt if pf.left_indent is not None else None,
    )


def table_to_rich(table: Table) -> RichTable:
    """Преобразует таблицу python-docx в :class:`RichTable`.

    Объединённые ячейки (gridSpan/vMerge) python-docx отдаёт повторами одного
    объекта по сетке — без учёта этого текст дублировался бы в каждой позиции
    («Итого» трижды в строке). Ячейка-исток получает colspan/rowspan, накрытые
    позиции помечаются ``merged`` и остаются пустыми.
    """
    tc_grid = [[cell._tc for cell in row.cells] for row in table.rows]
    nrows = len(tc_grid)
    rows: list[RichTableRow] = []
    for r, row in enumerate(table.rows):
        cells: list[RichTableCell] = []
        for c, cell in enumerate(row.cells):
            tc = tc_grid[r][c]
            if (c > 0 and tc_grid[r][c - 1] is tc) or (r > 0 and tc_grid[r - 1][c] is tc):
                cells.append(RichTableCell(merged=True))
                continue
            colspan = 1
            while c + colspan < len(tc_grid[r]) and tc_grid[r][c + colspan] is tc:
                colspan += 1
            rowspan = 1
            while (
                r + rowspan < nrows
                and c < len(tc_grid[r + rowspan])
                and tc_grid[r + rowspan][c] is tc
            ):
                rowspan += 1
            cells.append(
                RichTableCell(
                    paragraphs=[
                        para_to_rich(item)
                        for item in _iter_block_items(cell)
                        if isinstance(item, Paragraph)
                    ],
                    colspan=colspan,
                    rowspan=rowspan,
                )
            )
        rows.append(RichTableRow(cells=cells))
    return RichTable(rows=rows)


# --------------------------------------------------------------------------- #
#  Экстрактор
# --------------------------------------------------------------------------- #


class WordExtractor:
    """Чтение индекса, старых чисел и текстовых блоков из старого документа."""

    def _open(self, doc_path: Path) -> DocxDocument:
        try:
            return Document(str(doc_path))
        except Exception as exc:  # битый или чужой формат — в доменную ошибку
            raise DocumentParseError(f"Не удалось открыть '{doc_path.name}': {exc}") from exc

    # ------------------------------------------------------------------ #
    #  1. Индекс
    # ------------------------------------------------------------------ #
    def extract_index(self, doc_path: Path) -> str | None:
        doc = self._open(doc_path)
        for paragraph in doc.paragraphs[:25]:
            found = index_from_text(paragraph.text)
            if found:
                return found
        for table in doc.tables[:2]:
            for row in table.rows:
                for cell in row.cells:
                    found = index_from_text(cell.text)
                    if found:
                        return found
        return None

    # ------------------------------------------------------------------ #
    #  2. Старые числа (для diff)
    # ------------------------------------------------------------------ #
    def extract_old_numbers(self, doc_path: Path) -> OldNumbers:
        doc = self._open(doc_path)
        table = self._find_hours_table(doc)
        if table is None:
            return OldNumbers()
        return self._parse_hours_table(table)

    def _find_hours_table(self, doc: DocxDocument) -> Table | None:
        for table in doc.tables:
            header_text = " ".join(
                cell.text for row in table.rows[: min(3, len(table.rows))] for cell in row.cells
            )
            n = normalize_text(header_text)
            if "трудоемкость" in n and ("лекцион" in n or "практическ" in n):
                return table
        return None

    def _parse_hours_table(self, table: Table) -> OldNumbers:
        grid = [[cell.text for cell in row.cells] for row in table.rows]
        if not grid:
            return OldNumbers()
        ncols = max(len(r) for r in grid)

        def col_text(c: int) -> str:
            return " ".join(grid[r][c] for r in range(len(grid)) if c < len(grid[r]))

        col_labels = [normalize_text(col_text(c)) for c in range(ncols)]

        def find_col(*keywords: str) -> int | None:
            for c, label in enumerate(col_labels):
                if all(kw in label for kw in keywords):
                    return c
            for c, label in enumerate(col_labels):
                if any(kw in label for kw in keywords):
                    return c
            return None

        ctrud = find_col("трудоемкость")
        if ctrud is None:
            return OldNumbers()

        # строка данных — первая ниже шапки, где трудоёмкость числовая
        data_row: list[str] | None = None
        for r in range(len(grid)):
            if ctrud < len(grid[r]):
                hours, _ze = parse_hours_ze(grid[r][ctrud])
                if hours is not None and "трудоемкость" not in normalize_text(grid[r][ctrud]):
                    data_row = grid[r]
                    break
        if data_row is None:
            return OldNumbers()

        def val(c: int | None) -> str:
            return data_row[c] if c is not None and c < len(data_row) else ""

        hours_total, ze = parse_hours_ze(val(ctrud))
        semester, _course = parse_hours_ze(val(find_col("семестр")))

        # СРС может быть разбита на несколько колонок («Домашнее задание»,
        # «Самоконтроль» под общей шапкой «Виды самостоятельной работы») — суммируем.
        srs_vals = [
            _opt_int(val(c))
            for c, label in enumerate(col_labels)
            if "самостоятельн" in label or "срс" in label
        ]
        srs_known = [v for v in srs_vals if v is not None]
        # «Контроль, часов» — не путать с «Самоконтролем» из блока СРС.
        ctrl_col = next(
            (
                c
                for c, label in enumerate(col_labels)
                if "контрол" in label
                and "самоконтрол" not in label
                and "самостоятельн" not in label
            ),
            None,
        )
        return OldNumbers(
            ze=ze,
            hours_total=as_int(hours_total) if hours_total is not None else None,
            hours_lectures=_opt_int(val(find_col("лекцион"))),
            hours_practical=_opt_int(val(find_col("практическ"))),
            hours_lab=_opt_int(val(find_col("лабораторн"))),
            hours_project=_opt_int(val(find_col("проектн"))),
            hours_srs=sum(srs_known) if srs_known else None,
            hours_control=_opt_int(val(ctrl_col)),
            semester=int(semester) if semester is not None else None,
            # «форма» встречается во всех колонках из-за объединённой строки
            # «Очная форма обучения» — требуем и «аттестац» (точная колонка).
            control_raw=val(find_col("форма", "аттестац")) or None,
        )

    # ------------------------------------------------------------------ #
    #  3. Текстовые блоки (конечный автомат)
    # ------------------------------------------------------------------ #
    def extract(self, doc_path: Path, doc_type: DocType = DocType.RPD) -> ContentBlocks:
        doc = self._open(doc_path)
        result = ContentBlocks()
        self._fill_title_meta(doc, result)

        current: ContentBlock | None = None
        numbering = _NumberingModel(doc)
        # (numId, ilvl), рендерящиеся как «Задача № N: » — уровень обязателен:
        # варианты ответов часто живут на следующем уровне того же списка.
        task_levels: set[tuple[int, int]] = set()

        def switch_to(key: str | None, title: str) -> ContentBlock:
            if key is None or key == IGNORED_KEY:
                return ContentBlock(key=IGNORED_KEY, title=title)  # сток
            if key in result.blocks:
                return result.blocks[key]
            block = ContentBlock(key=key, title=title)
            result.blocks[key] = block
            return block

        for item in _iter_block_items(doc):
            if isinstance(item, Paragraph):
                text = item.text.strip()
                if _TOC_LINE_RE.search(text):
                    continue  # строка оглавления («…\t8») — не заголовок и не контент
                if doc_type == DocType.FOS and text and len(text) < 120:
                    # Маркеры типа вопросов (могут быть и ненумерованными):
                    # одобренные эталоны раскладывают задачи → §1, тесты → §2.
                    marker = next(
                        (
                            key
                            for key, kws in _FOS_MARKERS
                            if any(kw in normalize_text(text) for kw in kws)
                        ),
                        "",
                    )
                    if marker != "":
                        current = switch_to(marker, text)
                        continue
                if is_heading(text) or (
                    text
                    and len(text) <= _MAX_HEADING_LEN
                    and (_has_heading_style(item) or _is_bold_numbered_heading(item))
                ):
                    key = _heading_key(text)
                    if (
                        doc_type == DocType.FOS
                        and key is not None
                        and key not in _FOS_KEYS
                    ):
                        # Заголовок с РПД-ключом в ФОС (шкалы, критерии,
                        # «…компетенций») — его содержимому в ФОС 2026 не место
                        # (это источник §8 РПД): закрываем блок, контент в сток.
                        key = IGNORED_KEY
                    if key == IGNORED_KEY:
                        # Стоп-заголовок: закрывает текущий блок (иначе заголовки
                        # следующего раздела исходника утекают в перенесённый
                        # контент), а его собственный контент не сохраняется.
                        current = switch_to(None, text)
                        continue
                    if key is not None:
                        # Новый блок начинается ТОЛЬКО на распознанном заголовке;
                        # повторный ключ — продолжаем (merge) тот же блок, чтобы не
                        # терять контент при дублирующихся заголовках.
                        current = switch_to(key, text)
                        continue
                    # Нумерованный, но нераспознанный заголовок (например, пункт
                    # списка «1. Дайте определение…») — это контент текущего блока,
                    # а не новый раздел: не дробим списки вопросов на части.
                if current is not None and text:
                    if normalize_text(text).rstrip(".") in ("не предусмотрено", "не предусмотрены"):
                        # Заглушка исходника внутри блока — шум; для пустого
                        # блока генератор сам подставит «Не предусмотрено.».
                        continue
                    rich = para_to_rich(item)
                    props = _num_props(item)
                    if props is not None and rich.runs:
                        # Автонумерация Word при переносе теряется — синтезируем
                        # видимый номер/маркер в текст («1. », «– »).
                        task = False
                        if doc_type == DocType.FOS and current.key != IGNORED_KEY:
                            # Нумерованные вопросы в ФОС — формат эталона
                            # «Задача № N: …»; решение по первому абзацу уровня.
                            if (
                                props not in task_levels
                                and "?" in text
                                and numbering.is_decimal_dot(*props)
                            ):
                                task_levels.add(props)
                            task = props in task_levels
                        rich.runs[0].text = rich.runs[0].text.lstrip()
                        rich.runs.insert(0, RichRun(text=numbering.prefix(*props, task_style=task)))
                    current.elements.append(
                        ContentElement(kind=ElementKind.PARAGRAPH, paragraph=rich)
                    )
            elif isinstance(item, Table):
                if current is not None:
                    current.elements.append(
                        ContentElement(kind=ElementKind.TABLE, table=table_to_rich(item))
                    )
        return result

    def _fill_title_meta(self, doc: DocxDocument, result: ContentBlocks) -> None:
        paragraphs = [p.text.strip() for p in doc.paragraphs[:30]]
        for text in paragraphs:
            if not text:
                continue
            if result.index is None:
                idx = index_from_text(text)
                if idx:
                    result.index = idx
                    result.title_line = text
            low = normalize_text(text)
            if "по направлению подготовки" in low and result.direction is None:
                result.direction = re.sub(
                    r"(?i)по направлению подготовки\s*", "", text
                ).strip()
            elif low.startswith("направленности") and result.profile is None:
                result.profile = re.sub(r"(?i)направленности\s*\(профиля\)\s*", "", text).strip()
            elif "форма обучения" in low and result.form_study is None:
                result.form_study = re.sub(r"(?i)форма обучения\s*", "", text).strip()

        # Титульная таблица «метка | значение» (характерна для ФОС:
        # «Направление подготовки | 09.04.03 …») — абзацный скан её не видит.
        for table in doc.tables[:3]:
            for row in table.rows:
                if len(row.cells) < 2:
                    continue
                label = normalize_text(row.cells[0].text)
                value = " ".join(row.cells[1].text.split())
                if not value:
                    continue
                if "направление подготовки" in label and result.direction is None:
                    result.direction = value
                elif (
                    "образовательная программа" in label or "профил" in label
                ) and result.profile is None:
                    result.profile = value
                elif "форма обучения" in label and result.form_study is None:
                    result.form_study = value


def _opt_int(value: str) -> int | None:
    """parse_hours_ze → внешнее число (для ячеек вида «12» или «12 (…)»)."""
    outer, _ = parse_hours_ze(value)
    return as_int(outer) if outer is not None else None
