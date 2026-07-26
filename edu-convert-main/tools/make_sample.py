"""Генерация синтетического мини-примера в samples/ (без реальных данных вуза).

Создаёт самодостаточный демонстрационный набор, на котором можно прогнать
приложение, не имея реального учебного плана и документов:

* ``samples/план_пример.xlsx`` — мини учебный план с одной дисциплиной Б1.О.01;
* ``samples/Б1_О_01_РПД_Пример.docx`` — исходная РПД (со старыми числами,
  отличными от плана, чтобы продемонстрировать отчёт о расхождениях).

Запуск::

    python -m tools.make_sample
    python -m app.cli --db samples/план_пример.xlsx --input samples --out пример.zip
"""

from __future__ import annotations

from pathlib import Path

import xlsxwriter
from docx import Document

ROOT = Path(__file__).resolve().parent.parent
SAMPLES = ROOT / "samples"


def make_plan(path: Path) -> None:
    wb = xlsxwriter.Workbook(str(path))
    ws = wb.add_worksheet("План")
    ws.write(1, 11, "Семестр 1")  # заголовок блока семестра над колонками
    header = {
        0: "Считать в плане", 1: "Индекс", 2: "Наименование",
        3: "Экза мен", 4: "Зачет", 5: "Зачет с оц.",
        6: "Факт", 7: "По плану", 8: "Конт. раб.", 9: "СР", 10: "Конт роль",
        11: "з.е.", 12: "Лек", 13: "Лаб", 14: "Пр", 15: "Конс",
        16: "КаттЗ", 17: "ПО", 18: "СР", 19: "КаттЭ", 20: "Контроль",
        31: "Код", 32: "Компетенции",
    }
    for col, label in header.items():
        ws.write(2, col, label)
    discipline = {
        0: "+", 1: "Б1.О.01", 2: "Методология научных исследований (пример)",
        4: 1, 6: 3, 7: 108, 8: 30, 9: 78,
        11: 3, 12: 12, 14: 10, 17: 6, 18: 78,
        31: "01", 32: "УК-1-И-1; ПК-2-И-1",
    }
    for col, value in discipline.items():
        ws.write(4, col, value)
    wb.close()


def make_rpd(path: Path) -> None:
    doc = Document()
    doc.add_paragraph().add_run("Б1.О.01 МЕТОДОЛОГИЯ НАУЧНЫХ ИССЛЕДОВАНИЙ (ПРИМЕР)").bold = True
    doc.add_paragraph("по направлению подготовки 09.04.03 – «Прикладная информатика»")
    doc.add_paragraph("направленности (профиля) «Демонстрационный профиль»")
    doc.add_paragraph("форма обучения очная")

    doc.add_paragraph("1. Объем дисциплины и виды учебной работы")
    t = doc.add_table(rows=2, cols=5)
    for c, h in enumerate(["Семестр (курс)", "Форма промежуточной аттестации",
                           "Общая трудоемкость часов (ЗЕТ)", "Лекционные занятия, часов",
                           "Практические занятия, часов"]):
        t.rows[0].cells[c].text = h
    # Старые числа НАМЕРЕННО отличаются от плана (лек 16≠12, пр 24≠10) — для отчёта о расхождениях.
    for c, v in enumerate(["1 (1)", "Зачет", "108 (3)", "16", "24"]):
        t.rows[1].cells[c].text = v

    doc.add_paragraph("4. Учебно-методическое обеспечение дисциплины")
    doc.add_paragraph("4.1.1. Основная литература")
    p = doc.add_paragraph()
    p.add_run("Пример А.А. ").bold = True
    p.add_run("Демонстрационный источник. — М., 2024.")
    doc.add_paragraph("4.1.2. Дополнительная литература")
    doc.add_paragraph("Пример Б.Б. Дополнительный источник. — СПб., 2023.")
    doc.save(str(path))


def main() -> None:
    SAMPLES.mkdir(exist_ok=True)
    make_plan(SAMPLES / "план_пример.xlsx")
    make_rpd(SAMPLES / "Б1_О_01_РПД_Пример.docx")
    print(f"Синтетический пример создан в {SAMPLES}")


if __name__ == "__main__":
    main()
