"""Разметка официальных форм 2026 тегами docxtpl (этап B плана).

Берёт исходные пустые формы ГУУ (с редакторскими подписями) и порождает
размеченные шаблоны в ``app/templates/``:

* ``rpd_2026_tagged.docx`` — РПД;
* ``fos_2026_tagged.docx`` — ФОС.

Редакторские подписи («Из п. 2 Исходной РПД…», «Берём из учебного плана»)
заменяются тегами ``{{ }}``; пустые ячейки таблицы часов — тегами ``{{ hours_* }}``;
стандартные пред-заполненные разделы (ЭБС, ПО, помещения) остаются без изменений.

Скрипт идемпотентный и самодокументирующий: показывает, как именно получены
поставляемые шаблоны. Запуск::

    python -m tools.build_templates --rpd-src "Шаблон РП (2026) ММЭУ.docx" \
        --fos-src "ФОС_ШАБЛОН (2026) ММЭУ.docx"
"""

from __future__ import annotations

import argparse
import copy
from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph

from app.core.normalizer import normalize_text


def strip_red_highlights(doc: DocxDocument) -> None:
    """Удаляет красные highlight-пометки (в т.ч. на знаках абзаца), унаследованные
    из исходной формы 2026. Жёлтые (наши) не трогаем."""
    for h in list(doc.element.body.iter(qn("w:highlight"))):
        if (h.get(qn("w:val")) or "") == "red":
            h.getparent().remove(h)

PROJECT_ROOT = Path(__file__).resolve().parent.parent
TEMPLATES_DIR = PROJECT_ROOT / "app" / "templates"


# --------------------------------------------------------------------------- #
#  Утилиты редактирования docx
# --------------------------------------------------------------------------- #
def _capture_rpr(paragraph: Paragraph):
    """Снимок rPr первого содержательного run — чтобы при замене текста
    сохранить шрифт/размер исходной формы (титульные строки — явные 12pt).

    Цвет, заливка и подчёркивание исходного run не переносятся: редакторские
    подписи в формах красные, а подсветку/линии значений мы задаём сами."""
    runs = paragraph.runs
    rpr = None
    for r in runs:
        if r.text.strip() and r._element.rPr is not None:
            rpr = copy.deepcopy(r._element.rPr)
            break
    if rpr is None and runs and runs[0]._element.rPr is not None:
        rpr = copy.deepcopy(runs[0]._element.rPr)
    if rpr is not None:
        for tag in ("w:color", "w:highlight", "w:u"):
            el = rpr.find(qn(tag))
            if el is not None:
                rpr.remove(el)
    return rpr


def _apply_rpr(run, rpr) -> None:
    if rpr is None:
        return
    el = run._element
    if el.rPr is not None:
        el.remove(el.rPr)
    el.insert(0, copy.deepcopy(rpr))


def set_text(paragraph: Paragraph, text: str) -> None:
    """Заменяет весь текст абзаца одним runом (сохраняя стиль абзаца и rPr)."""
    rpr = _capture_rpr(paragraph)
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)
    _apply_rpr(paragraph.add_run(text), rpr)


def clear(paragraph: Paragraph) -> None:
    set_text(paragraph, "")


def hl_run(run) -> None:
    """Подсветка значений отключена по решению кафедры (раньше — жёлтая
    заливка всего заполненного); точки вызова сохранены."""
    return


def set_value_tag(paragraph: Paragraph, tag: str) -> None:
    """Абзац = только тег значения, жёлтый (рендеренное значение подсветится)."""
    set_text(paragraph, tag)
    hl_run(paragraph.runs[0])


def set_label_value(paragraph: Paragraph, label: str, tag: str) -> None:
    """Абзац = метка (без подсветки) + тег значения (жёлтый)."""
    rpr = _capture_rpr(paragraph)
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)
    _apply_rpr(paragraph.add_run(label), rpr)
    value = paragraph.add_run(tag)
    _apply_rpr(value, rpr)
    hl_run(value)


def _set_right_tab(paragraph: Paragraph, pos_twips: int) -> None:
    """Заменяет табстопы абзаца одним right-табом на ``pos_twips`` —
    хвостовой подчёркнутый таб дотягивает «нижнюю линию» ровно до него."""
    pPr = paragraph._p.get_or_add_pPr()
    old = pPr.find(qn("w:tabs"))
    if old is not None:
        pPr.remove(old)
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), str(pos_twips))
    tabs.append(tab)
    pPr.append(tabs)


def set_label_value_underlined(
    paragraph: Paragraph,
    label: str,
    tag: str,
    *,
    value_prefix: str = "",
    tail: tuple[str, ...] = ("\t",),
    right_tab_twips: int | None = None,
) -> None:
    """Титульная строка формы: метка + подчёркнутое значение (жёлтое) +
    подчёркнутые табы — «нижняя линия» до табстопов, как в исходной форме.

    ``value_prefix`` — пробельный отступ перед значением («\xa0\xa0»/пробел),
    ``tail`` — хвостовые run'ы линии; ``right_tab_twips`` — позиция правого
    табстопа линии (штатные стопы формы стоят на середине строки — линия
    обрывалась, до края её дотягивали только длинные значения).
    rPr исходных run (шрифт, размер 12pt) сохраняется — иначе строки становятся
    крупнее, переносятся и линии уезжают."""
    rpr = _capture_rpr(paragraph)
    if right_tab_twips is not None:
        _set_right_tab(paragraph, right_tab_twips)
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)
    _apply_rpr(paragraph.add_run(label), rpr)
    if value_prefix:
        pre = paragraph.add_run(value_prefix)
        _apply_rpr(pre, rpr)
        pre.underline = True
    value = paragraph.add_run(tag)
    _apply_rpr(value, rpr)
    value.underline = True
    hl_run(value)
    for tail_text in tail:
        t = paragraph.add_run(tail_text)
        _apply_rpr(t, rpr)
        t.underline = True


def insert_before(paragraph: Paragraph, text: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    np = Paragraph(new_p, paragraph._parent)
    np.add_run(text)
    return np


def remove_paragraph(paragraph: Paragraph) -> None:
    paragraph._element.getparent().remove(paragraph._element)


def insert_after(paragraph: Paragraph, text: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    np = Paragraph(new_p, paragraph._parent)
    np.add_run(text)
    return np


def n(text: str) -> str:
    return normalize_text(text)


# --------------------------------------------------------------------------- #
#  РПД
# --------------------------------------------------------------------------- #
def _style_letterhead(doc: DocxDocument) -> None:
    """Верхняя шапка титула («Федеральное государственное …», «„ГУУ“») —
    Tahoma 12 (требование кафедры; в форме размер не задан явно)."""
    for p in doc.paragraphs:
        if not p.text.strip():
            break
        for run in p.runs:
            run.font.name = "Tahoma"
            run.font.size = Pt(12)


def tag_rpd(src: Path, dst: Path) -> None:
    doc = Document(str(src))
    to_remove: list[Paragraph] = []

    # --- Абзацы: замена подписей на теги, очистка примеров --- #
    for p in doc.paragraphs:
        t = n(p.text)
        if not t:
            continue
        if "код и наименование дисциплины" in t and p.text.strip().isupper():
            set_value_tag(p, "{{ index }} {{ name }}")
        elif t.startswith("по направлению подготовки"):
            set_label_value_underlined(p, "по направлению подготовки ", "{{ direction }}",
                                       value_prefix=" ", right_tab_twips=9638)
        elif t.startswith("направленности"):
            set_label_value_underlined(p, "направленности (профиля) ", "{{ profile }}",
                                       right_tab_twips=9638)
        elif t.startswith("форма обучения"):
            set_label_value_underlined(p, "форма обучения ", "{{ form_study }}\xa0\xa0",
                                       value_prefix="\xa0\xa0", right_tab_twips=9638)
        elif "цели в исходной программе нет" in t:
            clear(p)  # целей в исходных РПД нет — раздел заполняется вручную
        # Subdoc-теги — с префиксом «p»: docxtpl убирает абзац-носитель и
        # вставляет блочный XML на уровень body (иначе w:p вкладывается в w:t —
        # невалидный OOXML, контент не виден python-docx и др. потребителям).
        elif "из п. 2 исходной рпд (столбцы 2,3)" in t:
            set_text(p, "{{p competencies }}")
        elif "из п. 2 исходной рпд (столбец 4)" in t:
            set_text(p, "{{p indicators }}")
        elif "из исходной рпд п.3" in t:
            # Тематический план не вставляется subdoc'ом: официальная таблица
            # тем заполняется генератором (часы масштабируются к Базе).
            to_remove.append(p)
        elif "указать основную литературу" in t:
            set_text(p, "{{p literature_main }}")
        elif "указать дополнительную литературу" in t:
            set_text(p, "{{p literature_extra }}")
        elif "указать ресурсы сети интернет" in t:
            set_text(p, "{{p internet_resources }}")
        elif "берем из учебного плана" in t:
            clear(p)
        elif t.startswith("заседания кафедры"):
            set_label_value_underlined(p, "заседания кафедры ", "{{ department_name }}",
                                       value_prefix="\xa0\xa0", tail=("\t ",))
        elif "методов в экономике и управлении" in t:
            to_remove.append(p)  # хвост захардкоженного названия кафедры (вторая строка)
        elif t.startswith("опк-1"):  # примеры компетенций/индикаторов
            clear(p)
        elif t == "тематический план дисциплины":
            # В эталоне тематический план начинается с новой страницы.
            p.paragraph_format.page_break_before = True
    for p in to_remove:
        remove_paragraph(p)

    _style_letterhead(doc)

    # --- Таблица часов: значения из Excel --- #
    _tag_hours_table(doc)

    # --- Остальные редакторские пометки (§3 тематический план, §8 оценивание) --- #
    _clean_editorial(doc)
    _tidy_rpd_blanks(doc)
    _tighten_title_before_moscow(doc)
    _fix_fos_page_break(doc)  # «1. Цели…» начинает страницу свойством, не прокладкой
    strip_red_highlights(doc)

    dst.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(dst))
    print(f"[РПД] {src.name} → {dst}")


def _tighten_title_before_moscow(doc: DocxDocument) -> None:
    """Убирает 2 пустых абзаца перед «Москва, …»: реальные значения титула
    (название, профиль) длиннее заглушек формы и «Москва» уезжала со страницы."""
    paragraphs = doc.paragraphs
    for i, p in enumerate(paragraphs):
        if n(p.text).startswith("москва"):
            removed = 0
            for q in reversed(paragraphs[:i]):
                if removed == 2:
                    break
                if not q.text.strip():
                    remove_paragraph(q)
                    removed += 1
                else:
                    break
            break


def _tidy_rpd_blanks(doc: DocxDocument) -> None:
    """Подгоняет пустые абзацы под эталонную вёрстку.

    Эталон держит ровно один пустой абзац после блоков §2 (компетенции,
    индикаторы) и три пустых перед названием дисциплины на титуле."""
    paragraphs = doc.paragraphs

    def trim_after(idx: int, keep: int) -> None:
        for seen, q in enumerate(paragraphs[idx + 1 :], start=1):
            if q.text.strip():
                break
            if seen > keep:
                remove_paragraph(q)

    for i, p in enumerate(paragraphs):
        t = p.text.strip()
        if t in ("{{p competencies }}", "{{p indicators }}"):
            trim_after(i, keep=1)
        elif t == "{{ index }} {{ name }}":
            # перед названием — три пустых: реальные значения титула длиннее
            # однострочных заглушек эталона, иначе «Москва, 2026» уезжает.
            blanks = []
            for q in reversed(paragraphs[:i]):
                if q.text.strip():
                    break
                blanks.append(q)
            for q in blanks[3:]:
                remove_paragraph(q)


def _clean_editorial(doc: DocxDocument) -> None:
    """Убирает/заполняет оставшиеся редакторские пометки в РПД.

    * §3: убирает инструкции «Указать несколько тем…» и «Проверка итого…»
      (тематический план заполняется автором по официальной пустой таблице);
    * §8: примечание «Исходная РПД п.п.8…» убирается, ячейки «Код/коды
      компетенций» → тег ``{{ competence_codes }}`` (из Базы), инструкции
      «(перечислить)» удаляются.
    """
    # Верхнеуровневые красные примечания.
    for p in doc.paragraphs:
        t = n(p.text)
        if "проверка итого" in t or "исходная рпд п. п. 8" in t or "столбцы 2 и 5" in t:
            clear(p)

    grade_level = {
        "отлично": "5", "хорошо": "4", "удовлетворительно": "3", "неудовлетворительно": "2",
    }
    for table in doc.tables:
        header = n(" ".join(c.text for c in table.rows[0].cells))
        if "формируемая компетенция" in header:  # §8 «Система оценивания»
            for row in table.rows:
                level = grade_level.get(n(row.cells[0].text))
                if level is None:
                    continue  # строка-шапка
                # c1 = «Формируемая компетенция» → родительские коды (жёлтые)
                if len(row.cells) > 1:
                    set_value_tag(row.cells[1].paragraphs[0], "{{ competence_parents }}")
                    for extra in row.cells[1].paragraphs[1:]:
                        clear(extra)
                # c2 = «Наименование результатов»: оставить вводную фразу (подсветить),
                # убрать «Знает/Умеет/Владеет (перечислить)», добавить тег результатов.
                if len(row.cells) > 2:
                    cell = row.cells[2]
                    for p in list(cell.paragraphs):
                        if "демонстрирует" in n(p.text):
                            for run in p.runs:
                                if run.text.strip():
                                    hl_run(run)
                        else:  # удаляем «Знает/Умеет/Владеет (перечислить)» без пустых строк
                            p._element.getparent().remove(p._element)
                    hl_run(cell.add_paragraph().add_run("{{ outcomes_" + level + " }}"))
        else:  # §3 «Тематический план» и прочие
            for row in table.rows:
                for cell in row.cells:
                    if "указать несколько тем" in n(cell.text):
                        for p in cell.paragraphs:
                            clear(p)


def _tag_hours_table(doc: DocxDocument) -> None:
    # строка-метка (нормализованная) → имя тега для колонки «Всего» (c2) и «семестр» (c3)
    row_tags = {
        "зач. ед.": "{{ ze }}",
        "ак.ч.": "{{ hours_total }}",
        "часы контактной работы": "{{ hours_contact }}",
        "из них аудиторной работы": "{{ hours_aud }}",
        "лекции": "{{ hours_lectures }}",
        "практические занятия": "{{ hours_practical }}",
        "лабораторные занятия": "{{ hours_lab }}",
        "проектное обучение": "{{ hours_project }}",
        "часы внеаудиторной работы": "{{ hours_extra_contact }}",
        "часы самостоятельной работы": "{{ hours_self_study }}",
        "вид промежуточной аттестации": "{{ control_summary }}",
    }
    for table in doc.tables:
        header = " ".join(c.text for c in table.rows[0].cells)
        if "Вид учебной работы" not in header:
            continue
        # Вторая строка шапки (vMerge-продолжение) и пустые прокладки:
        # в эталонах их нет. Текст vMerge-ячейки видится из истока, поэтому
        # проверяем собственное содержимое w:tc.
        for row in list(table.rows):
            own_texts = [
                "".join(t.text or "" for t in tc.iter(qn("w:t"))).strip()
                for tc in row._tr.iter(qn("w:tc"))
            ]
            if not any(own_texts):
                row._tr.getparent().remove(row._tr)
        # «4 семестр» в шапке колонки «Кол-во часов в семестре (ак.ч.)».
        header_row = table.rows[0]
        if len(header_row.cells) > 3:
            sem_cell = header_row.cells[3]
            rpr = _capture_rpr(sem_cell.paragraphs[0])
            sem_p = sem_cell.add_paragraph()
            sem_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _apply_rpr(sem_p.add_run("{{ semester_label }}"), rpr)
        for row in table.rows:
            # «Общая трудоёмкость» в c0, а единица («зач. ед.»/«ак.ч.») — в c1
            label = n(row.cells[0].text + " " + (row.cells[1].text if len(row.cells) > 1 else ""))
            tag = next((v for k, v in row_tags.items() if k in label), None)
            if tag is None:
                continue
            # колонка «Всего» = индекс 2; «в семестре» = 3 (для односеместровых равны).
            # Значения подсвечиваются жёлтым (всё, что заполнено из Базы).
            # Аттестация — вид строчными в обеих ячейках («экзамен | экзамен»).
            cell_tag = "{{ control_kind }}" if "{{ control" in tag else tag
            if len(row.cells) > 2:
                set_value_tag(row.cells[2].paragraphs[0], cell_tag)
            if len(row.cells) > 3:
                set_value_tag(row.cells[3].paragraphs[0], cell_tag)
        break


# --------------------------------------------------------------------------- #
#  ФОС
# --------------------------------------------------------------------------- #
def tag_fos(src: Path, dst: Path) -> None:
    doc = Document(str(src))
    gia_heading: Paragraph | None = None
    to_remove: list[Paragraph] = []

    for p in doc.paragraphs:
        t = n(p.text)
        if not t:
            continue
        if t == "код наименование":
            set_value_tag(p, "{{ index }} {{ name }}")
        elif t.startswith("по направлению подготовки"):
            # Пробельные отступы перед значениями — как в форме/эталонах.
            set_label_value_underlined(p, "по направлению подготовки ", "{{ direction }}",
                                       value_prefix=" ", right_tab_twips=9638)
        elif t.startswith("направленности"):
            set_label_value_underlined(p, "направленности (профиля) ", "{{ profile }}",
                                       right_tab_twips=9638)
        elif t.startswith("форма обучения"):
            set_label_value_underlined(p, "форма обучения ", "{{ form_study }}\xa0\xa0",
                                       value_prefix="\xa0\xa0", right_tab_twips=9638)
        # Статики формы «Примерный перечень задач» / «Примерный состав тестовых
        # вопросов…» сохраняются (так в одобренных эталонах) — контент после них.
        elif "примерный перечень задач" in t:
            pass  # статика формы, контент вставляется после «Задачи к разделу 1…»
        elif "примерный состав тестовых вопросов" in t:
            insert_after(p, "{{p interim_attestation }}")
        elif t.startswith("задачи к разделу 1"):
            # Первая строка формы остаётся: индикаторы — из Базы (жёлтым),
            # разбивку по разделам уточняет методист. Контент задач — после неё.
            set_label_value(p, "Задачи к разделу 1. (оцениваемая компетенция и индикатор ",
                            "{{ fos_indicators }}")
            p.add_run(")")
            insert_after(p, "{{p current_control }}")
        elif t.startswith("задачи к разделу"):
            to_remove.append(p)  # вторая строка-пример формы
        elif t.startswith("задача") or t == "ответ:":
            to_remove.append(p)  # примеры формы: удаляем абзац целиком, не оставляя пустот
        elif "обязательно с ответами" in t:
            to_remove.append(p)
        elif t.startswith("заседания кафедры"):
            set_label_value_underlined(p, "заседания кафедры ", "{{ department_name }}",
                                       value_prefix="\xa0\xa0", tail=("\t ",))
        elif "методов в экономике и управлении" in t:
            to_remove.append(p)  # хвост захардкоженного названия кафедры (вторая строка)
        elif "государственной итоговой" in t:
            gia_heading = p

    _style_letterhead(doc)
    if gia_heading is not None:
        # Раздела 3 нет вовсе, если дисциплина не участвует в ГИА (исходный ФОС
        # без блока ГИА): заголовок + контент под jinja-условием.
        insert_before(gia_heading, "{%p if gia_present %}")
        tag_p = insert_after(gia_heading, "{{p gia }}")
        insert_after(tag_p, "{%p endif %}")
    for p in to_remove:
        remove_paragraph(p)
    _fix_fos_page_break(doc)

    strip_red_highlights(doc)
    dst.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(dst))
    print(f"[ФОС] {src.name} → {dst}")


def _fix_fos_page_break(doc: DocxDocument) -> None:
    """Раздел 1 начинает страницу свойством pageBreakBefore.

    В форме разрыв страницы — отдельный пустой абзац после «Москва, 2026»:
    при длинных значениях титула (название, профиль) он уезжал на следующую
    страницу и оставлял пустой лист. Абзацы-прокладки удаляются, разрыв
    становится свойством первого содержательного абзаца."""
    paragraphs = doc.paragraphs
    mi = next((i for i, p in enumerate(paragraphs) if n(p.text).startswith("москва")), None)
    if mi is None:
        return
    for p in paragraphs[mi + 1 :]:
        if p.text.strip():
            p.paragraph_format.page_break_before = True
            break
        remove_paragraph(p)



def main() -> None:
    ap = argparse.ArgumentParser(description="Разметка шаблонов 2026 тегами docxtpl")
    _src_dir = PROJECT_ROOT / "files" if (PROJECT_ROOT / "files").is_dir() else PROJECT_ROOT
    ap.add_argument("--rpd-src", default=str(_src_dir / "Шаблон РП (2026) ММЭУ.docx"))
    ap.add_argument("--fos-src", default=str(_src_dir / "ФОС_ШАБЛОН (2026) ММЭУ.docx"))
    args = ap.parse_args()

    tag_rpd(Path(args.rpd_src), TEMPLATES_DIR / "rpd_2026_tagged.docx")
    tag_fos(Path(args.fos_src), TEMPLATES_DIR / "fos_2026_tagged.docx")


if __name__ == "__main__":
    main()
