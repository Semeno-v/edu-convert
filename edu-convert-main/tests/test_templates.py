"""Регрессионные проверки поставляемых размеченных шаблонов.

Гарантируют, что в шаблонах не осталось редакторских пометок (красный текст,
инструкции «Указать…», «(перечислить)», «Из п. 2 Исходной РПД» и т.п.), которые
иначе попадут в сгенерированные документы.
"""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document
from docx.oxml.ns import qn
from docxtpl import DocxTemplate

from app.config import settings

_EDITORIAL_PHRASES = [
    "указать несколько тем",
    "проверка итого",
    "исходная рпд",
    "(перечислить)",
    "коды компетенций",
    "из п. 2 исходной",
    "берем из учебного плана",
    "цели в исходной программе нет",
]


def _all_text(path: Path) -> str:
    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts).lower()


def _red_runs(path: Path) -> list[str]:
    doc = Document(str(path))

    def is_red(run: object) -> bool:
        rpr = run._element.find(qn("w:rPr"))
        if rpr is None:
            return False
        col = rpr.find(qn("w:color"))
        if col is None:
            return False
        c = col.get(qn("w:val")) or ""
        if c in ("auto", "000000"):
            return False
        try:
            r, g, b = int(c[0:2], 16), int(c[2:4], 16), int(c[4:6], 16)
        except ValueError:
            return False
        return r > 120 and r > g + 40 and r > b + 40

    paragraphs = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)
    return [r.text.strip() for p in paragraphs for r in p.runs if is_red(r) and r.text.strip()]


@pytest.fixture(params=["rpd", "fos"])
def template_path(request: pytest.FixtureRequest) -> Path:
    return settings.rpd_template if request.param == "rpd" else settings.fos_template


def test_template_exists(template_path: Path) -> None:
    assert template_path.exists(), f"шаблон не найден: {template_path}"


def test_template_no_red_marks(template_path: Path) -> None:
    reds = _red_runs(template_path)
    assert reds == [], f"в шаблоне остались красные пометки: {reds}"


def test_template_no_editorial_text(template_path: Path) -> None:
    text = _all_text(template_path)
    leftover = [phrase for phrase in _EDITORIAL_PHRASES if phrase in text]
    assert leftover == [], f"в шаблоне остались редакторские подписи: {leftover}"


def test_rpd_template_has_expected_tags() -> None:
    # thematic_plan тегом не является: официальную таблицу тем заполняет
    # генератор после рендера (часы масштабируются к Базе).
    tags = DocxTemplate(str(settings.rpd_template)).get_undeclared_template_variables()
    for tag in (
        "competence_parents", "competencies", "indicators", "control_kind", "hours_self_study",
    ):
        assert tag in tags, f"в шаблоне РПД нет тега {tag}"


def test_template_no_yellow_highlight(template_path: Path) -> None:
    # Подсветка отключена по решению кафедры — в шаблонах её быть не должно.
    import zipfile

    xml = zipfile.ZipFile(template_path).read("word/document.xml").decode("utf-8", "replace")
    assert '<w:highlight w:val="yellow"' not in xml


def test_rpd_template_letterhead_tahoma12() -> None:
    # Верхняя шапка титула — Tahoma 12 (требование кафедры).
    doc = Document(str(settings.rpd_template))
    first = doc.paragraphs[0]
    assert "Федеральное" in first.text
    for run in first.runs:
        assert run.font.name == "Tahoma"
        assert run.font.size is not None and round(run.font.size.pt) == 12


def test_rpd_template_semester_label_in_hours_header() -> None:
    # Семестр прописывается в шапке колонки «Кол-во часов в семестре».
    doc = Document(str(settings.rpd_template))
    hours = next(t for t in doc.tables if "Вид учебной работы" in t.rows[0].cells[0].text)
    assert "{{ semester_label }}" in hours.rows[0].cells[3].text
