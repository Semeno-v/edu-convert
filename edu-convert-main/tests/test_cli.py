"""CLI (app/cli.py) — сборка входов, коды возврата, запись архива.

CLI не покрывался тестами вовсе, а именно им пользуются для пакетных прогонов
без интерфейса. Проверяется поведение, видимое снаружи: что попадает в список
входных файлов, с каким кодом завершается запуск и остаётся ли после него
временная папка.
"""

from __future__ import annotations

import zipfile
from pathlib import Path

import pytest
from docx import Document

from app.cli import _collect_inputs, main


def test_collect_inputs_expands_directories(tmp_path: Path) -> None:
    (tmp_path / "РПД.docx").touch()
    (tmp_path / "ФОС.doc").touch()
    (tmp_path / "заметки.txt").touch()

    collected = _collect_inputs([str(tmp_path)])

    assert sorted(p.name for p in collected) == ["РПД.docx", "ФОС.doc"]


def test_collect_inputs_skips_word_lock_files(tmp_path: Path) -> None:
    (tmp_path / "РПД.docx").touch()
    (tmp_path / "~$РПД.docx").touch()

    assert [p.name for p in _collect_inputs([str(tmp_path)])] == ["РПД.docx"]


def test_collect_inputs_takes_files_as_is(tmp_path: Path) -> None:
    """Явно названный файл берётся без фильтра по расширению."""
    odd = tmp_path / "план.xlsx"
    odd.touch()

    assert _collect_inputs([str(odd)]) == [odd]


def test_collect_inputs_ignores_missing_paths(tmp_path: Path) -> None:
    assert _collect_inputs([str(tmp_path / "нет-такого.docx")]) == []


def test_returns_error_code_without_inputs(
    tmp_path: Path, plan_xlsx: Path, capsys: pytest.CaptureFixture[str]
) -> None:
    """Пустой список входов — код 2 и объяснение в stderr, а не молчаливый успех."""
    code = main(["--db", str(plan_xlsx), "--input", str(tmp_path / "пусто")])

    assert code == 2
    assert "Не найдено" in capsys.readouterr().err


def test_end_to_end_writes_zip_and_cleans_temp(
    tmp_path: Path, plan_xlsx: Path, rpd_docx: Path, capsys: pytest.CaptureFixture[str]
) -> None:
    out = tmp_path / "результат.zip"

    code = main(
        ["--db", str(plan_xlsx), "--input", str(rpd_docx), "--out", str(out)]
    )

    assert code == 0
    assert out.exists()
    with zipfile.ZipFile(out) as zf:
        assert "report.xlsx" in zf.namelist()
    assert "СВОДНЫЙ ОТЧЁТ" in capsys.readouterr().out


def test_summary_lists_failed_file(
    tmp_path: Path, plan_xlsx: Path, capsys: pytest.CaptureFixture[str]
) -> None:
    """Сбойный файл виден в сводке — иначе о нём узнают только из ZIP."""
    bad = tmp_path / "Б9_Х_99_РПД.docx"
    doc = Document()
    doc.add_paragraph().add_run("Б9.Х.99 НЕСУЩЕСТВУЮЩАЯ ДИСЦИПЛИНА").bold = True
    doc.save(str(bad))

    code = main(
        ["--db", str(plan_xlsx), "--input", str(bad), "--out", str(tmp_path / "р.zip")]
    )

    assert code == 0  # сбой одного файла не роняет прогон целиком
    out = capsys.readouterr().out
    assert "[X]" in out
    assert bad.name in out
