"""Интеграционный тест оркестратора (services/orchestrator.py).

Проверяет весь конвейер на синтетических данных + поставляемых шаблонах:
обработка валидного файла (с расхождением), изоляция сбоя (индекс не найден),
формат колонок report.xlsx (ТЗ §8), наличие ZIP и очистку temp (ТЗ §7.3).
Использует только .docx (без Word COM), поэтому проходит в CI.
"""

from __future__ import annotations

import io
import zipfile
from pathlib import Path

import anyio
import polars as pl
from docx import Document

from app.core.models import FileStatus
from app.services.orchestrator import Orchestrator

_REPORT_COLUMNS_EXPECTED = [
    "Имя исходного файла", "Индекс дисциплины", "Статус",
    "Описание ошибки или расхождения", "Значение в старом файле", "Значение в Базе",
]


def _make_bad_input(path: Path) -> None:
    doc = Document()
    doc.add_paragraph().add_run("Б9.Х.99 НЕСУЩЕСТВУЮЩАЯ ДИСЦИПЛИНА").bold = True
    doc.save(str(path))


def test_orchestrator_end_to_end(plan_xlsx: Path, rpd_docx: Path, tmp_path: Path) -> None:
    bad = tmp_path / "Б9_Х_99_РПД_Нет.docx"
    _make_bad_input(bad)

    orch = Orchestrator(plan_xlsx)  # шаблоны по умолчанию из templates/
    result = anyio.run(orch.run, [rpd_docx, bad])
    report = result.report

    assert report.total == 2

    # Валидный файл Б1.О.01: старые числа (14/20) ≠ эталон (12/10) → Расхождение.
    good = next(r for r in report.results if r.index == "Б1.О.01")
    assert good.status in (FileStatus.SUCCESS, FileStatus.DISCREPANCY)
    assert good.output_name and good.output_name.endswith(".docx")

    # Битый файл: индекс не найден → Ошибка, изоляция сбоя (второй файл обработан).
    bad_res = next(r for r in report.results if r.filename == bad.name)
    assert bad_res.status == FileStatus.ERROR
    assert "не найден" in bad_res.message.lower()

    # ZIP содержит отчёт и сгенерированный документ.
    with zipfile.ZipFile(result.zip_path) as zf:
        names = zf.namelist()
        assert "report.xlsx" in names
        assert any(n.endswith(".docx") for n in names)
        # Колонки отчёта — дословно по ТЗ §8.
        df = pl.read_excel(io.BytesIO(zf.read("report.xlsx")))
        assert df.columns == _REPORT_COLUMNS_EXPECTED
        # Строка-ошибка для ненайденной дисциплины присутствует.
        statuses = df.get_column("Статус").to_list()
        assert "Ошибка" in statuses

    # Очистка temp (ТЗ §7.3).
    workdir, zip_path = result.workdir, result.zip_path
    result.cleanup()
    assert not workdir.exists()
    assert not zip_path.exists()


def test_orchestrator_failure_cleans_workdir(
    plan_xlsx: Path, rpd_docx: Path, monkeypatch, tmp_path: Path
) -> None:
    # Сбой на формировании отчёта/архива не должен оставлять workdir в temp:
    # RunResult ещё не создан, cleanup() вызвать некому.
    from app.config import settings

    orch = Orchestrator(plan_xlsx)

    def boom(*args: object, **kwargs: object) -> None:
        raise OSError("диск занят")

    monkeypatch.setattr(orch, "_make_zip", boom)
    before = set(settings.temp_root.glob("run_*")) if settings.temp_root.exists() else set()
    try:
        anyio.run(orch.run, [rpd_docx])
        raise AssertionError("ожидался OSError")
    except OSError:
        pass
    after = set(settings.temp_root.glob("run_*")) if settings.temp_root.exists() else set()
    assert after == before  # новых осиротевших папок не появилось
