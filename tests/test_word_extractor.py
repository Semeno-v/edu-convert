"""Тесты экстрактора Word (цель ≥95% покрытия, ТЗ §7.5).

Покрывают распознавание заголовков, конечный автомат блоков, преобразование в
IR (с сохранением форматирования), разбор старых чисел и извлечение индекса.
"""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from app.core.exceptions import DocumentParseError
from app.core.models import DocType, ElementKind
from app.core.word_extractor import (
    WordExtractor,
    _heading_key,
    _opt_int,
    is_heading,
    para_to_rich,
    table_to_rich,
)


@pytest.fixture
def extractor() -> WordExtractor:
    return WordExtractor()


# --------------------------------------------------------------------------- #
#  Распознавание и классификация заголовков
# --------------------------------------------------------------------------- #
@pytest.mark.parametrize(
    ("text", "expected"),
    [
        ("1. Объем дисциплины", True),
        ("3.Формирование компетентностной траектории", True),  # без пробела после номера
        ("4.1.1. Основная литература", True),
        ("8.1. Шкала оценивания", True),
        ("обычный абзац без номера", False),
        ("", False),
        ("1 2 3 4 5 6 7", True),  # начинается с числа — формально заголовок
        ("X" * 250, False),  # слишком длинный
    ],
)
def test_is_heading(text: str, expected: bool) -> None:
    assert is_heading(text) is expected


@pytest.mark.parametrize(
    ("text", "expected_key"),
    [
        ("2. Роль дисциплины в формировании компетенций", "competencies"),
        ("4.1.1. Основная литература", "literature_main"),
        ("4.1.2. Дополнительная литература", "literature_extra"),
        ("6. Программное обеспечение", "software"),
        ("5. Материально-техническое обеспечение", "facilities"),
        ("1.1 Контрольные вопросы для текущего контроля", "current_control"),
        ("3. Вопросы для государственной итоговой аттестации", "gia"),
        ("Цели освоения дисциплины", "goals"),
        ("Нечто непонятное про погоду", None),
    ],
)
def test_heading_key(text: str, expected_key: str | None) -> None:
    assert _heading_key(text) == expected_key




# --------------------------------------------------------------------------- #
#  Индекс
# --------------------------------------------------------------------------- #
def test_extract_index(extractor: WordExtractor, rpd_docx: Path) -> None:
    assert extractor.extract_index(rpd_docx) == "Б1.О.01"


def test_extract_index_from_table(extractor: WordExtractor, tmp_path: Path) -> None:
    doc = Document()
    t = doc.add_table(rows=1, cols=1)
    t.rows[0].cells[0].text = "Дисциплина Б1.В.05 Инфраструктура"
    path = tmp_path / "notitle.docx"
    doc.save(str(path))
    assert extractor.extract_index(path) == "Б1.В.05"


def test_extract_index_none(extractor: WordExtractor, empty_docx: Path) -> None:
    assert extractor.extract_index(empty_docx) is None


# --------------------------------------------------------------------------- #
#  Старые числа (для diff)
# --------------------------------------------------------------------------- #
def test_extract_old_numbers(extractor: WordExtractor, rpd_docx: Path) -> None:
    old = extractor.extract_old_numbers(rpd_docx)
    assert old.ze == 3.0
    assert old.hours_total == 108
    assert old.hours_lectures == 14
    assert old.hours_practical == 20
    assert old.semester == 1
    assert old.control_raw is not None and "Зачет" in old.control_raw


def test_extract_old_numbers_no_table(extractor: WordExtractor, empty_docx: Path) -> None:
    old = extractor.extract_old_numbers(empty_docx)
    assert old.ze is None
    assert old.hours_total is None


def test_extract_old_numbers_real_layout(extractor: WordExtractor, tmp_path: Path) -> None:
    # Реальная шапка: строка «Очная форма обучения» во всех колонках (бывшее
    # объединение) загрязняет label каждой колонки словом «форма»; СРС разбита
    # на «Домашнее задание» и «Самоконтроль» под общей шапкой.
    doc = Document()
    hdr = [
        "Семестр (курс)", "Форма промежуточной аттестации",
        "Общая трудоемкость, часов (ЗЕТ)", "Лекционные занятия, часов",
        "Практические занятия, часов",
        "Виды самостоятельной работы – Домашнее задание, часов",
        "Виды самостоятельной работы – Самоконтроль, часов",
        "Контроль, часов",
    ]
    data = ["4(2)", "Экзамен", "216 (6)", "16", "32", "90", "31", "30"]
    t = doc.add_table(rows=3, cols=len(hdr))
    for c, h in enumerate(hdr):
        t.rows[0].cells[c].text = h
        t.rows[1].cells[c].text = "Очная форма обучения"
        t.rows[2].cells[c].text = data[c]
    path = tmp_path / "real_hours.docx"
    doc.save(str(path))

    old = extractor.extract_old_numbers(path)
    assert old.control_raw == "Экзамен"      # не колонка семестра
    assert old.hours_srs == 121              # 90 + 31
    assert old.hours_control == 30
    assert old.semester == 4
    assert old.ze == 6.0 and old.hours_total == 216


# --------------------------------------------------------------------------- #
#  Конечный автомат: текстовые блоки
# --------------------------------------------------------------------------- #
def test_extract_content_blocks(extractor: WordExtractor, rpd_docx: Path) -> None:
    content = extractor.extract(rpd_docx, DocType.RPD)
    assert content.index == "Б1.О.01"
    assert content.direction is not None and "09.04.03" in content.direction
    assert content.profile is not None
    assert content.form_study is not None and "очная" in content.form_study

    assert "competencies" in content.blocks
    assert "literature_main" in content.blocks
    assert "literature_extra" in content.blocks
    assert "software" in content.blocks


def test_competencies_block_has_table(extractor: WordExtractor, rpd_docx: Path) -> None:
    content = extractor.extract(rpd_docx, DocType.RPD)
    comp = content.get("competencies")
    assert comp is not None
    kinds = [e.kind for e in comp.elements]
    assert ElementKind.TABLE in kinds


def test_literature_preserves_bold(extractor: WordExtractor, rpd_docx: Path) -> None:
    content = extractor.extract(rpd_docx, DocType.RPD)
    main = content.get("literature_main")
    assert main is not None and not main.is_empty
    para = main.elements[0].paragraph
    assert para is not None
    assert any(run.bold for run in para.runs)  # «Иванов И.И.» жирным


def test_extract_empty_doc(extractor: WordExtractor, empty_docx: Path) -> None:
    content = extractor.extract(empty_docx, DocType.RPD)
    assert content.blocks == {}
    assert content.index is None


# --------------------------------------------------------------------------- #
#  Раны внутри контейнеров (w:hyperlink, w:smartTag) — paragraph.runs их теряет
# --------------------------------------------------------------------------- #
def _wrap_run(container_tag: str, text: str):
    container = OxmlElement(container_tag)
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    r.append(t)
    container.append(r)
    return container


def test_para_to_rich_keeps_hyperlink_text() -> None:
    doc = Document()
    p = doc.add_paragraph("URL: ")
    p._p.append(_wrap_run("w:hyperlink", "https://e.lib/1"))
    rich = para_to_rich(p)
    assert rich.text == "URL: https://e.lib/1"


def test_para_to_rich_keeps_smarttag_text() -> None:
    doc = Document()
    p = doc.add_paragraph("поля: нижнее ")
    p._p.append(_wrap_run("w:smartTag", "2 см"))
    rich = para_to_rich(p)
    assert "2 см" in rich.text


def test_para_to_rich_hyperlink_only_paragraph() -> None:
    doc = Document()
    p = doc.add_paragraph()
    p._p.append(_wrap_run("w:hyperlink", "www.sovnet.ru"))
    rich = para_to_rich(p)
    assert rich.text == "www.sovnet.ru"


# --------------------------------------------------------------------------- #
#  Стоп-заголовки: хвосты следующего раздела не утекают в перенесённый блок
# --------------------------------------------------------------------------- #
def test_stop_headings_close_block(extractor: WordExtractor, tmp_path: Path) -> None:
    doc = Document()
    doc.add_paragraph("3. Структура и содержание дисциплины")
    doc.add_paragraph("Тема 1. Введение")
    doc.add_paragraph("4. Учебно-методическое обеспечение дисциплины")
    doc.add_paragraph("4.1 Рекомендуемая литература по дисциплине")
    doc.add_paragraph("4.1.1. Основная литература")
    doc.add_paragraph("Иванов И.И. Книга.")
    path = tmp_path / "stop.docx"
    doc.save(str(path))

    content = extractor.extract(path, DocType.RPD)
    plan = content.get("thematic_plan")
    assert plan is not None
    plan_text = " ".join(e.paragraph.text for e in plan.elements if e.paragraph)
    assert "Тема 1" in plan_text
    assert "Учебно-методическое" not in plan_text   # хвост не утёк в §3
    assert "Рекомендуемая литература" not in plan_text
    main = content.get("literature_main")
    assert main is not None and not main.is_empty   # литература ловится своим правилом


def test_fos_drops_assessment_scales(extractor: WordExtractor, tmp_path: Path) -> None:
    # Одобренные эталоны ФОС 2026 не содержат шкал/таблиц оценивания (их место —
    # §8 РПД): «Описание шкал…» закрывает блок, контент шкал отбрасывается.
    doc = Document()
    doc.add_paragraph("1.2 Контрольные вопросы для промежуточной аттестации")
    doc.add_paragraph("Содержательный вопрос без знака.")
    doc.add_paragraph("2. Описание шкал оценивания степени сформированности компетенций")
    doc.add_paragraph("Зачтено — навыки сформированы.")
    t = doc.add_table(rows=1, cols=2)
    t.cell(0, 0).text = "Оценка"
    t.cell(0, 1).text = "Отлично"
    path = tmp_path / "fos_scales.docx"
    doc.save(str(path))

    content = extractor.extract(path, DocType.FOS)
    block = content.get("interim_attestation")
    assert block is not None
    text = " ".join(e.paragraph.text for e in block.elements if e.paragraph)
    assert "Содержательный вопрос" in text
    assert "Описание шкал" not in text                    # заголовок не переносится
    assert "навыки сформированы" not in text              # контент шкал отброшен
    assert all(e.table is None for e in block.elements)   # таблица шкал не попала
    assert "competencies" not in content.blocks

    # В РПД РПД-ключи работают как раньше («компетенц» → свой блок).
    rpd = extractor.extract(path, DocType.RPD)
    assert "competencies" in rpd.blocks


def test_bold_numbered_list_headings(extractor: WordExtractor, tmp_path: Path) -> None:
    # Заголовки-списки («List Paragraph» + numPr + целиком жирные) — как в РПД
    # факультативов; обычные нумерованные абзацы контента не дробят блок.
    doc = Document()
    h = doc.add_paragraph("Основная литература")
    h.runs[0].bold = True
    _numbered(doc, "")  # numPr добавим вручную ниже на заголовок
    # numPr на жирный заголовок
    numPr = OxmlElement("w:numPr")
    nid = OxmlElement("w:numId"); nid.set(qn("w:val"), "7")
    numPr.append(nid)
    h._p.get_or_add_pPr().append(numPr)
    doc.add_paragraph("Иванов И.И. Книга про язык.")
    path = tmp_path / "boldnum.docx"
    doc.save(str(path))

    content = extractor.extract(path, DocType.RPD)
    main = content.get("literature_main")
    assert main is not None
    assert any(e.paragraph and "Иванов" in e.paragraph.text for e in main.elements)


def test_styled_headings_without_numbers(extractor: WordExtractor, tmp_path: Path) -> None:
    # Заголовки без видимых номеров (автонумерация в стиле на базе Heading)
    # распознаются по цепочке стилей — иначе блоки таких РПД не извлекались.
    doc = Document()
    styles = doc.styles
    from docx.enum.style import WD_STYLE_TYPE

    h = styles.add_style("Кастомный заголовок", WD_STYLE_TYPE.PARAGRAPH)
    h.base_style = styles["Heading 2"]
    p1 = doc.add_paragraph("Основная литература"); p1.style = h
    doc.add_paragraph("Иванов И.И. Книга.")
    p2 = doc.add_paragraph("Дополнительная литература"); p2.style = h
    doc.add_paragraph("Петров П.П. Другая книга.")
    path = tmp_path / "styled.docx"
    doc.save(str(path))

    content = extractor.extract(path, DocType.RPD)
    main, extra = content.get("literature_main"), content.get("literature_extra")
    assert main is not None and "Иванов" in main.elements[0].paragraph.text
    assert extra is not None and "Петров" in extra.elements[0].paragraph.text


def test_title_meta_from_title_table(extractor: WordExtractor, tmp_path: Path) -> None:
    # Титул ФОС: направление/программа/форма лежат в таблице «метка | значение».
    doc = Document()
    doc.add_paragraph("Б1.О.01 Методология научных исследований")
    t = doc.add_table(rows=3, cols=2)
    t.cell(0, 0).text = "Направление подготовки"
    t.cell(0, 1).text = "09.04.03 «Прикладная информатика»"
    t.cell(1, 0).text = "Образовательная программа"
    t.cell(1, 1).text = "Цифровые технологии в управлении"
    t.cell(2, 0).text = "Форма обучения"
    t.cell(2, 1).text = "очная"
    path = tmp_path / "fos_title.docx"
    doc.save(str(path))

    content = extractor.extract(path, DocType.FOS)
    assert content.direction == "09.04.03 «Прикладная информатика»"
    assert content.profile == "Цифровые технологии в управлении"
    assert content.form_study == "очная"


# --------------------------------------------------------------------------- #
#  Синтез номеров автонумерации (numbering.xml теряется при переносе)
# --------------------------------------------------------------------------- #
def _numbered(doc, text: str, num_id: int = 5, ilvl: int = 0):
    p = doc.add_paragraph(text)
    numPr = OxmlElement("w:numPr")
    lvl = OxmlElement("w:ilvl")
    lvl.set(qn("w:val"), str(ilvl))
    nid = OxmlElement("w:numId")
    nid.set(qn("w:val"), str(num_id))
    numPr.append(lvl)
    numPr.append(nid)
    p._p.get_or_add_pPr().append(numPr)
    return p


def test_extract_synthesizes_list_numbers(extractor: WordExtractor, tmp_path: Path) -> None:
    # В ФОС нумерованные вопросы получают формат одобренных эталонов
    # («Задача № N: …»); в РПД — обычный синтез номера («N. …»).
    doc = Document()
    doc.add_paragraph("1.1 Контрольные вопросы для текущего контроля")
    _numbered(doc, "Что такое наука?")
    _numbered(doc, "Что такое знание?")
    path = tmp_path / "numbered.docx"
    doc.save(str(path))

    content = extractor.extract(path, DocType.FOS)
    block = content.get("current_control")
    assert block is not None
    texts = [e.paragraph.text for e in block.elements if e.paragraph]
    assert texts == ["Задача № 1: Что такое наука?", "Задача № 2: Что такое знание?"]

    rpd = extractor.extract(path, DocType.RPD)
    rpd_block = rpd.get("current_control")
    assert rpd_block is not None
    rpd_texts = [e.paragraph.text for e in rpd_block.elements if e.paragraph]
    assert rpd_texts == ["1. Что такое наука?", "2. Что такое знание?"]


def test_numbering_model_formats() -> None:
    from app.core.word_extractor import _NumberingModel

    nm = _NumberingModel.__new__(_NumberingModel)
    nm._levels = {
        (7, 0): ("bullet", "", 1),
        (8, 0): ("decimal", "%1)", 1),
        (9, 0): ("lowerLetter", "%1)", 1),
    }
    nm._counters = {}
    assert nm.prefix(7, 0) == "– "                      # маркер
    # варианты «N)» в эталонах отделяются табом
    assert [nm.prefix(8, 0) for _ in range(2)] == ["1)\t", "2)\t"]
    assert nm.prefix(9, 0) == "a)\t"                    # lowerLetter
    assert nm.prefix(99, 0) == "1. "                    # неизвестный numId — decimal


# --------------------------------------------------------------------------- #
#  Объединённые ячейки: текст не дублируется, размах сохраняется
# --------------------------------------------------------------------------- #
def test_table_to_rich_preserves_merges(tmp_path: Path) -> None:
    doc = Document()
    t = doc.add_table(rows=3, cols=3)
    t.cell(0, 0).merge(t.cell(0, 2))      # горизонталь: вся первая строка
    t.cell(1, 0).merge(t.cell(2, 0))      # вертикаль: первая колонка вниз
    t.cell(0, 0).text = "Шапка"
    t.cell(1, 0).text = "Итого"
    t.cell(1, 1).text = "10"

    rich = table_to_rich(t)
    head = rich.rows[0].cells
    assert head[0].colspan == 3 and not head[0].merged
    assert head[1].merged and head[2].merged
    assert head[1].paragraphs == []        # текст «Шапка» не задвоен
    col = [rich.rows[1].cells[0], rich.rows[2].cells[0]]
    assert col[0].rowspan == 2 and not col[0].merged
    assert col[1].merged

    texts = [
        p.text
        for row in rich.rows
        for cell in row.cells
        for p in cell.paragraphs
    ]
    assert texts.count("Шапка") == 1
    assert texts.count("Итого") == 1


def test_stop_heading_content_dropped(extractor: WordExtractor, tmp_path: Path) -> None:
    doc = Document()
    doc.add_paragraph("4.2 Ресурсы сети Интернет")
    doc.add_paragraph("https://example.org")
    doc.add_paragraph("4.3 Дополнительные средства обучения (в том числе on-line курсы)")
    doc.add_paragraph("Не предусмотрено.")
    doc.add_paragraph("4.4 Профессиональные базы данных и информационно-справочные системы")
    doc.add_paragraph("СПС Гарант")
    path = tmp_path / "drop.docx"
    doc.save(str(path))

    content = extractor.extract(path, DocType.RPD)
    res = content.get("internet_resources")
    assert res is not None
    text = " ".join(e.paragraph.text for e in res.elements if e.paragraph)
    assert "https://example.org" in text
    assert "СПС Гарант" in text            # ПБД сливаются в интернет-ресурсы
    assert "Не предусмотрено" not in text  # заглушка-абзац отфильтрована
    assert "Дополнительные средства" not in text  # заголовок поглощён merge'м


def test_additional_means_table_preserved(extractor: WordExtractor, tmp_path: Path) -> None:
    # Таблица под «4.3 Дополнительные средства обучения» (например, «Совнет») —
    # реальные on-line ресурсы, переносится в блок интернет-ресурсов.
    doc = Document()
    doc.add_paragraph("4.2 Ресурсы сети Интернет")
    doc.add_paragraph("https://example.org")
    doc.add_paragraph("4.3 Дополнительные средства обучения (в том числе on-line курсы)")
    t = doc.add_table(rows=1, cols=2)
    t.cell(0, 0).text = "Российская ассоциация управления проектами («Совнет»)"
    t.cell(0, 1).text = "www.sovnet.ru"
    path = tmp_path / "sovnet.docx"
    doc.save(str(path))

    content = extractor.extract(path, DocType.RPD)
    res = content.get("internet_resources")
    assert res is not None
    cell_text = " ".join(
        p.text
        for e in res.elements
        if e.table
        for row in e.table.rows
        for cell in row.cells
        for p in cell.paragraphs
    )
    assert "Совнет" in cell_text and "www.sovnet.ru" in cell_text


# --------------------------------------------------------------------------- #
#  IR-преобразования напрямую
# --------------------------------------------------------------------------- #
def test_para_to_rich_runs(tmp_path: Path) -> None:
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("жирный").bold = True
    p.add_run(" обычный")
    rich = para_to_rich(p)
    assert rich.text == "жирный обычный"
    assert rich.runs[0].bold is True
    assert rich.runs[1].bold is False
    assert rich.list_level is None


def test_para_to_rich_captures_style(tmp_path: Path) -> None:
    doc = Document()
    p = doc.add_paragraph("элемент списка", style="List Bullet")
    rich = para_to_rich(p)
    assert rich.style == "List Bullet"


def test_list_level_with_numpr() -> None:
    doc = Document()
    p = doc.add_paragraph("элемент")
    pPr = p._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "2")
    numPr.append(ilvl)
    pPr.append(numPr)
    assert para_to_rich(p).list_level == 2


def test_table_to_rich() -> None:
    doc = Document()
    t = doc.add_table(rows=2, cols=2)
    t.rows[0].cells[0].text = "A"
    t.rows[0].cells[1].text = "B"
    t.rows[1].cells[0].text = "C"
    rich = table_to_rich(t)
    assert len(rich.rows) == 2
    assert rich.rows[0].cells[0].paragraphs[0].text == "A"


def test_open_corrupt_file(extractor: WordExtractor, tmp_path: Path) -> None:
    bad = tmp_path / "bad.docx"
    bad.write_bytes(b"not a real docx")
    with pytest.raises(DocumentParseError):
        extractor.extract_index(bad)


def test_old_numbers_header_only(extractor: WordExtractor, tmp_path: Path) -> None:
    # Таблица с нужной шапкой, но без числовой строки данных → пустой OldNumbers.
    doc = Document()
    doc.add_paragraph("1. Объем дисциплины")
    t = doc.add_table(rows=1, cols=3)
    t.rows[0].cells[0].text = "Общая трудоёмкость часов (ЗЕТ)"
    t.rows[0].cells[1].text = "Лекционные занятия"
    t.rows[0].cells[2].text = "Практические занятия"
    path = tmp_path / "header_only.docx"
    doc.save(str(path))
    old = extractor.extract_old_numbers(path)
    assert old.hours_total is None


@pytest.mark.parametrize(
    ("value", "expected"),
    [("", None), ("12", 12), ("18 (5)", 18)],
)
def test_opt_int(value: str, expected: int | None) -> None:
    assert _opt_int(value) == expected


def test_list_level_numpr_without_ilvl() -> None:
    doc = Document()
    p = doc.add_paragraph("текст")
    pPr = p._p.get_or_add_pPr()
    pPr.append(OxmlElement("w:numPr"))  # numPr без ilvl → уровень 0
    assert para_to_rich(p).list_level == 0


def test_parse_hours_table_without_trudoemkost(extractor: WordExtractor) -> None:
    doc = Document()
    t = doc.add_table(rows=2, cols=2)
    t.rows[0].cells[0].text = "Колонка"
    t.rows[0].cells[1].text = "Другая"
    t.rows[1].cells[0].text = "1"
    # столбца «трудоёмкость» нет → пустой результат
    assert extractor._parse_hours_table(t).hours_total is None


def test_heading_merge_same_key(extractor: WordExtractor, tmp_path: Path) -> None:
    # Повторный распознанный заголовок СЛИВАЕТСЯ в тот же блок (а не дублируется),
    # чтобы не терять контент при дублирующихся/оглавлениях-заголовках.
    doc = Document()
    doc.add_paragraph("6. Программное обеспечение")
    doc.add_paragraph("Первое ПО")
    doc.add_paragraph("7. Программное обеспечение второй раз")
    doc.add_paragraph("Второе ПО")
    path = tmp_path / "dup.docx"
    doc.save(str(path))
    content = extractor.extract(path, DocType.RPD)
    assert "software" in content.blocks
    assert not any(k.startswith("software_") for k in content.blocks)
    texts = [e.paragraph.text for e in content.blocks["software"].elements if e.paragraph]
    assert "Первое ПО" in texts and "Второе ПО" in texts


def test_unrecognized_numbered_heading_is_content(extractor: WordExtractor, tmp_path: Path) -> None:
    # Нумерованный пункт списка (не раздел) остаётся контентом текущего блока.
    doc = Document()
    doc.add_paragraph("6. Программное обеспечение")
    doc.add_paragraph("1. Какой-то вопрос списка?")
    doc.add_paragraph("2. Ещё вопрос списка?")
    path = tmp_path / "list.docx"
    doc.save(str(path))
    content = extractor.extract(path, DocType.RPD)
    texts = [e.paragraph.text for e in content.blocks["software"].elements if e.paragraph]
    assert "1. Какой-то вопрос списка?" in texts and "2. Ещё вопрос списка?" in texts
    assert len(content.blocks) == 1  # вопросы не создали новых блоков
