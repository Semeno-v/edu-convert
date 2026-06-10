"""Тесты генератора docxtpl (core/word_generator.py).

Покрывают валидацию шаблона (обе ветки), рендеринг с числами из Excel
(золотое правило), вставку текстового блока и форматирование чисел.
Используются поставляемые размеченные шаблоны templates/*.docx.
"""

from __future__ import annotations

import zipfile
from pathlib import Path

import pytest
from docx import Document

from app.config import settings
from app.core.exceptions import TemplateValidationError
from app.core.models import (
    CompetencyGroup,
    CompetencyIndicator,
    ContentBlock,
    ContentBlocks,
    ContentElement,
    ControlForm,
    ControlKind,
    DocType,
    ElementKind,
    RichParagraph,
    RichRun,
    RichTable,
    RichTableCell,
    RichTableRow,
    SubjectData,
)
from app.core.word_generator import DocxtplGenerator, _assessment_outcomes, _num


def _assessment_block() -> ContentBlock:
    def cell(t: str) -> RichTableCell:
        return RichTableCell(paragraphs=[RichParagraph(runs=[RichRun(text=t)])])
    rows = [
        RichTableRow(cells=[cell("Оценка"), cell("Компетенция"), cell("Формулировка требований")]),
        RichTableRow(cells=[cell("Отлично"), cell("ПК-1"), cell("Знает на высоком уровне X")]),
        # vMerge-продолжение уровня «Отлично»: уровень наследуется
        RichTableRow(cells=[cell(""), cell("ПК-1"), cell("Знает также Y")]),
        RichTableRow(cells=[cell(""), cell("ПК-1"), cell("Умеет применять X")]),
        RichTableRow(cells=[cell("Неудовлетворительно"), cell("ПК-1"), cell("Не знает X")]),
    ]
    return ContentBlock(key="assessment", title="Оценка качества",
                        elements=[ContentElement(kind=ElementKind.TABLE, table=RichTable(rows=rows))])


@pytest.fixture
def generator() -> DocxtplGenerator:
    return DocxtplGenerator()


@pytest.fixture
def subject() -> SubjectData:
    return SubjectData(
        index="Б1.О.01", name="Тестовая дисциплина", ze=3.0, hours_total=108,
        hours_contact=30, hours_aud=28, hours_lectures=12, hours_practical=10,
        hours_lab=0, hours_project=6, hours_srs=78, hours_control=0,
        control_forms=(ControlForm(kind=ControlKind.CREDIT, semester=1),),
        competence_codes=("УК-1-И-1", "ПК-2-И-1"),
        competencies=(
            CompetencyGroup(code="УК-1", text="Способен критически мыслить",
                            indicators=(CompetencyIndicator(code="УК-1-И-1", text="Анализирует проблему"),)),
            CompetencyGroup(code="ПК-2", text="Способен моделировать",
                            indicators=(CompetencyIndicator(code="ПК-2-И-1", text="Строит модель"),)),
        ),
    )


def _lit_table() -> RichTable:
    def cell(t: str) -> RichTableCell:
        return RichTableCell(paragraphs=[RichParagraph(runs=[RichRun(text=t)])])
    head = RichTableRow(cells=[cell("№"), cell("Автор(ы)"), cell("Наименование"),
                               cell("Выходные данные"), cell("URL")])
    row = RichTableRow(cells=[cell("1"), cell("Иванов И.И."), cell("Книга про науку"),
                              cell("М., 2023. 200 с."), cell("URL: https://e.lib/1")])
    return RichTable(rows=[head, row])


@pytest.fixture
def content() -> ContentBlocks:
    block = ContentBlock(
        key="literature_main", title="Основная литература",
        elements=[ContentElement(kind=ElementKind.TABLE, table=_lit_table())],
    )
    return ContentBlocks(index="Б1.О.01", direction="09.04.03", profile="Профиль",
                         form_study="очная", blocks={"literature_main": block})


@pytest.mark.parametrize(("value", "expected"), [(3.0, 3), (2.5, 2.5), (6, 6), (4.0, 4)])
def test_num_formatting(value: float, expected: object) -> None:
    assert _num(value) == expected


def test_validate_template_ok(generator: DocxtplGenerator) -> None:
    # На поставляемых размеченных шаблонах валидация проходит без исключений.
    generator.validate_template(settings.rpd_template, DocType.RPD)
    generator.validate_template(settings.fos_template, DocType.FOS)


def test_validate_template_missing_tags(generator: DocxtplGenerator, tmp_path: Path) -> None:
    doc = Document()
    doc.add_paragraph("Документ без обязательных тегов")
    bad = tmp_path / "bad_template.docx"
    doc.save(str(bad))
    with pytest.raises(TemplateValidationError):
        generator.validate_template(bad, DocType.RPD)


def test_generate_uses_excel_numbers(
    generator: DocxtplGenerator, subject: SubjectData, content: ContentBlocks, tmp_path: Path
) -> None:
    out = tmp_path / "out.docx"
    generator.generate(settings.rpd_template, out, subject, content, DocType.RPD)
    assert out.exists()

    doc = Document(str(out))
    hours: dict[str, str] = {}
    for table in doc.tables:
        if "Вид учебной работы" in " ".join(c.text for c in table.rows[0].cells):
            for row in table.rows:
                label = row.cells[0].text.strip()
                val = row.cells[2].text.strip() if len(row.cells) > 2 else ""
                if label == "Лекции":
                    hours["lec"] = val
                elif label == "Практические занятия":
                    hours["pr"] = val
    # Числа строго из Excel (SubjectData), не из Word.
    assert hours.get("lec") == "12"
    assert hours.get("pr") == "10"

    # Аттестация: вид строчными в обеих ячейках («экзамен | экзамен» в эталоне).
    for table in doc.tables:
        if "Вид учебной работы" in " ".join(c.text for c in table.rows[0].cells):
            att = next(r for r in table.rows if "аттестации" in r.cells[0].text)
            assert att.cells[2].text.strip() == "зачет"
            assert att.cells[3].text.strip() == "зачет"
            # «Часы самостоятельной работы» = чистая СРС из Базы (без контроля)
            srs = next(r for r in table.rows if "самостоятельной" in r.cells[0].text)
            assert srs.cells[2].text.strip() == "78"

    xml = zipfile.ZipFile(out).read("word/document.xml").decode("utf-8", "replace")
    # Литература переформатирована из таблицы в список — автор присутствует.
    assert "Иванов" in xml
    # Жёлтая подсветка заполненного присутствует.
    assert '<w:highlight w:val="yellow"' in xml


def test_competencies_built_from_base(
    generator: DocxtplGenerator, subject: SubjectData, tmp_path: Path
) -> None:
    # §2 строится из Базы (subject.competencies), даже при пустом ContentBlocks.
    out = tmp_path / "out_comp.docx"
    generator.generate(settings.rpd_template, out, subject, ContentBlocks(), DocType.RPD)
    xml = zipfile.ZipFile(out).read("word/document.xml").decode("utf-8", "replace")
    assert "Способен критически мыслить" in xml  # текст компетенции
    assert "Анализирует проблему" in xml          # текст индикатора
    assert "ПК-2" in xml                           # родительский код в §8


def test_assessment_outcomes_credit_grade() -> None:
    # Уровень «Зачет» (без «зачтено») распознаётся как лучший; «Не зачтено» — нет.
    def cell(t: str) -> RichTableCell:
        return RichTableCell(paragraphs=[RichParagraph(runs=[RichRun(text=t)])])
    rows = [
        RichTableRow(cells=[cell("Оценка"), cell("Формулировка")]),
        RichTableRow(cells=[cell("Зачет"), cell("Обучающийся умеет X.")]),
        RichTableRow(cells=[cell(""), cell("Обучающийся знает Y.")]),  # vMerge
        RichTableRow(cells=[cell("Не зачтено"), cell("Не умеет X.")]),
    ]
    block = ContentBlock(key="assessment", title="8.1",
                         elements=[ContentElement(kind=ElementKind.TABLE, table=RichTable(rows=rows))])
    out = _assessment_outcomes(block)
    assert out["5"] == "Обучающийся умеет X\aОбучающийся знает Y"
    assert "Не умеет" not in out["2"]


def test_internet_resources_filter_ebs(
    generator: DocxtplGenerator, subject: SubjectData, tmp_path: Path
) -> None:
    # §4.3: дубли ЭБС из §5 отфильтровываются, настоящие ресурсы — списком.
    def cell(t: str) -> RichTableCell:
        return RichTableCell(paragraphs=[RichParagraph(runs=[RichRun(text=t)])])
    table = RichTable(rows=[
        RichTableRow(cells=[cell("№"), cell("Наименование"), cell("Адрес доступа"), cell("Доступ")]),
        RichTableRow(cells=[cell("1"), cell("ЭБС Znanium"), cell("https://znanium.com"), cell("Договор")]),
        RichTableRow(cells=[cell("2"), cell("Команда в MS Teams «Курс»"), cell("https://teams.microsoft.com"), cell("Бесплатно")]),
    ])
    cb = ContentBlocks(blocks={"internet_resources": ContentBlock(
        key="internet_resources", title="4.2",
        elements=[ContentElement(kind=ElementKind.TABLE, table=table)])})
    out = tmp_path / "res.docx"
    generator.generate(settings.rpd_template, out, subject, cb, DocType.RPD)
    doc = Document(str(out))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "1. Команда в MS Teams «Курс». https://teams.microsoft.com. Бесплатно." in text
    assert "znanium" not in text.lower()


def test_assessment_outcomes_by_grade() -> None:
    # Формат эталонов: формулировки лучшего уровня группируются по глаголу
    # («Знает …; …» / «Умеет …») и тиражируются на все уровни.
    out = _assessment_outcomes(_assessment_block())
    expected = "Знает на высоком уровне X; также Y\aУмеет применять X"
    assert out["5"] == expected
    assert out["2"] == expected  # та же формулировка, вводную даёт шаблон


def test_section8_outcomes_filled(
    generator: DocxtplGenerator, subject: SubjectData, tmp_path: Path
) -> None:
    # §8 «Наименование результатов» заполняется из §8 исходной РПД:
    # формулировки лучшего уровня — во всех четырёх строках.
    cb = ContentBlocks(blocks={"assessment": _assessment_block()})
    out = tmp_path / "out8.docx"
    generator.generate(settings.rpd_template, out, subject, cb, DocType.RPD)
    xml = zipfile.ZipFile(out).read("word/document.xml").decode("utf-8", "replace")
    assert xml.count("Знает на высоком уровне X") == 4  # на каждом уровне
    assert "Умеет применять X" in xml
    assert "Не знает X" not in xml  # формулировки худших уровней не используются


def test_generate_empty_block_placeholder(
    generator: DocxtplGenerator, subject: SubjectData, tmp_path: Path
) -> None:
    # Пустой контент -> subdoc с заглушкой, без ошибки рендеринга.
    out = tmp_path / "out_empty.docx"
    generator.generate(settings.rpd_template, out, subject, ContentBlocks(index="Б1.О.01"),
                       DocType.RPD)
    assert out.exists()


def _topics_src_table() -> RichTable:
    """Таблица §3 исходника: шапка с видами занятий + 2 темы + итого."""
    def cell(t: str) -> RichTableCell:
        return RichTableCell(paragraphs=[RichParagraph(runs=[RichRun(text=t)])])

    return RichTable(rows=[
        RichTableRow(cells=[cell("Этапы"), cell("Дескрипторы"), cell("Содержание"),
                            cell("Лекционные занятия"), cell("Практические занятия"),
                            cell("Проектное обучение"), cell("Итого")]),
        RichTableRow(cells=[cell(""), cell("Д-1"), cell("Тема-АБВ"),
                            cell("3"), cell("2,5"), cell(""), cell("5,5")]),
        RichTableRow(cells=[cell(""), cell("Д-2"), cell("Тема-ГДЕ"),
                            cell("3"), cell("2,5"), cell("1"), cell("6,5")]),
        RichTableRow(cells=[cell("Итого по дисциплине:"), cell(""), cell(""),
                            cell("6"), cell("5"), cell("1"), cell("12")]),
    ])


def test_generate_fills_official_topics_table(
    generator: DocxtplGenerator, subject: SubjectData, tmp_path: Path
) -> None:
    # Официальная таблица тем заполняется из исходника, часы масштабируются
    # к Базе: лекции 6 → 12 (×2), практические 5 → 10 (×2), проектное 1 → 6.
    cb = ContentBlocks(blocks={"thematic_plan": ContentBlock(
        key="thematic_plan", title="ТП",
        elements=[ContentElement(kind=ElementKind.TABLE, table=_topics_src_table())])})
    out = tmp_path / "out_them.docx"
    generator.generate(settings.rpd_template, out, subject, cb, DocType.RPD)

    doc = Document(str(out))
    topics = next(t for t in doc.tables if "Темы (разделы)" in t.rows[0].cells[1].text)
    rows = [[c.text.strip() for c in r.cells] for r in topics.rows]
    tema = next(r for r in rows if r[1] == "Тема-АБВ")
    assert tema[0] == "1"
    assert tema[2] == "6"      # лекции 3 × (12/6)
    assert tema[3] == "5"      # практические 2,5 × (10/5)
    assert tema[4] == "-"      # лабораторных нет ни в исходнике, ни в Базе
    assert tema[5] == "-"      # проектное: пусто в исходнике → «-»
    assert tema[6] == "11"     # итого по строке
    itogo = next(r for r in rows if r[1] == "ИТОГО")
    assert itogo[2:] == ["12", "10", "-", "6", "28"]  # числа Базы (aud=28)
    # заполненные ячейки подсвечены
    xml = zipfile.ZipFile(out).read("word/document.xml").decode("utf-8", "replace")
    assert "Тема-АБВ" in xml and '<w:highlight w:val="yellow"' in xml


def test_generate_fos(
    generator: DocxtplGenerator, subject: SubjectData, content: ContentBlocks, tmp_path: Path
) -> None:
    out = tmp_path / "out_fos.docx"
    generator.generate(settings.fos_template, out, subject, content, DocType.FOS)
    assert out.exists()
    xml = zipfile.ZipFile(out).read("word/document.xml").decode("utf-8", "replace")
    assert "Б1.О.01" in xml  # индекс на титуле


def test_fos_body_format_and_indicators(
    generator: DocxtplGenerator, subject: SubjectData, tmp_path: Path
) -> None:
    # Контент ФОС — в эталонном формате тела (по ширине, красная строка 1 см);
    # строка «Задачи к разделу 1…» содержит индикаторы из Базы.
    block = ContentBlock(
        key="current_control", title="1.1",
        elements=[ContentElement(
            kind=ElementKind.PARAGRAPH,
            paragraph=RichParagraph(runs=[RichRun(text="Задача № 1: Что такое X?")]),
        )],
    )
    cb = ContentBlocks(index="Б1.О.01", blocks={"current_control": block})
    out = tmp_path / "fos_fmt.docx"
    generator.generate(settings.fos_template, out, subject, cb, DocType.FOS)

    doc = Document(str(out))
    target = next(p for p in doc.paragraphs if "Что такое X?" in p.text)
    assert target.paragraph_format.alignment is not None
    assert "JUSTIFY" in str(target.paragraph_format.alignment)
    assert round(target.paragraph_format.first_line_indent.pt, 1) == 28.4
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Задачи к разделу 1. (оцениваемая компетенция и индикатор УК-1-И-1, ПК-2-И-1)" in text


def _nested_block_in_wt(xml: str) -> list[str]:
    """Имена дочерних элементов внутри текстовых узлов w:t (валидный OOXML — пусто)."""
    from lxml import etree

    w = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    root = etree.fromstring(xml.encode("utf-8"))
    return [child.tag for t in root.iter(w + "t") for child in t]


@pytest.mark.parametrize("doc_type", [DocType.RPD, DocType.FOS])
def test_subdoc_content_is_valid_ooxml(
    generator: DocxtplGenerator, subject: SubjectData, content: ContentBlocks,
    tmp_path: Path, doc_type: DocType,
) -> None:
    # Subdoc-теги размечены как {{p …}}: блочный XML ложится на уровень body,
    # а не внутрь w:t (это невалидный OOXML, его не видит python-docx).
    template = settings.rpd_template if doc_type == DocType.RPD else settings.fos_template
    out = tmp_path / f"valid_{doc_type.value}.docx"
    generator.generate(template, out, subject, content, doc_type)
    xml = zipfile.ZipFile(out).read("word/document.xml").decode("utf-8", "replace")
    assert _nested_block_in_wt(xml) == []


def test_generate_table_restores_merges(
    generator: DocxtplGenerator, subject: SubjectData, tmp_path: Path
) -> None:
    # Объединённые ячейки исходника восстанавливаются (gridSpan/vMerge),
    # текст истока не дублируется по сетке.
    def cell(t: str = "", **kw: object) -> RichTableCell:
        paras = [RichParagraph(runs=[RichRun(text=t)])] if t else []
        return RichTableCell(paragraphs=paras, **kw)

    table = RichTable(rows=[
        RichTableRow(cells=[cell("Шапка-АБВ", colspan=3), cell(merged=True), cell(merged=True)]),
        RichTableRow(cells=[cell("Итого-XYZ", rowspan=2), cell("10"), cell("20")]),
        RichTableRow(cells=[cell(merged=True), cell("1"), cell("2")]),
    ])
    cb = ContentBlocks(blocks={"current_control": ContentBlock(
        key="current_control", title="КВ",
        elements=[ContentElement(kind=ElementKind.TABLE, table=table)])})
    out = tmp_path / "merged.docx"
    generator.generate(settings.fos_template, out, subject, cb, DocType.FOS)
    xml = zipfile.ZipFile(out).read("word/document.xml").decode("utf-8", "replace")
    assert xml.count("Шапка-АБВ") == 1
    assert xml.count("Итого-XYZ") == 1
    assert '<w:gridSpan w:val="3"/>' in xml
    assert "<w:vMerge" in xml


def test_generate_large_table_fills_all_cells(
    generator: DocxtplGenerator, subject: SubjectData, tmp_path: Path
) -> None:
    # Регрессия: id()-кэш lxml-прокси ложно помечал ячейки заполненными —
    # большие перенесённые таблицы рендерились с пустыми хвостами.
    rows = [
        RichTableRow(cells=[
            RichTableCell(paragraphs=[RichParagraph(runs=[RichRun(text=f"яч-{r}-{c}")])])
            for c in range(4)
        ])
        for r in range(40)
    ]
    cb = ContentBlocks(blocks={"current_control": ContentBlock(
        key="current_control", title="КВ",
        elements=[ContentElement(kind=ElementKind.TABLE, table=RichTable(rows=rows))])})
    out = tmp_path / "big.docx"
    generator.generate(settings.fos_template, out, subject, cb, DocType.FOS)
    xml = zipfile.ZipFile(out).read("word/document.xml").decode("utf-8", "replace")
    missing = [f"яч-{r}-{c}" for r in range(40) for c in range(4) if f"яч-{r}-{c}" not in xml]
    assert missing == []


def test_subdoc_content_visible_to_python_docx(
    generator: DocxtplGenerator, subject: SubjectData, content: ContentBlocks, tmp_path: Path
) -> None:
    # После валидной вставки subdoc-контент виден обычным python-docx.
    out = tmp_path / "visible.docx"
    generator.generate(settings.rpd_template, out, subject, content, DocType.RPD)
    doc = Document(str(out))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Способен критически мыслить" in text  # §2 из Базы
    assert "Иванов" in text                        # §4 литература из исходника
