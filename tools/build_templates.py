"""Разметка официальных форм 2026 тегами docxtpl (этап B плана).

Берёт исходные пустые формы ГУУ (с редакторскими подписями) и порождает
размеченные шаблоны в ``templates/``:

* ``rpd_2026_tagged.docx`` — РПД;
* ``fos_2026_tagged.docx`` — ФОС.

Редакторские подписи («Из п. 2 Исходной РПД…», «Берём из учебного плана»)
заменяются тегами ``{{ }}``; пустые ячейки таблицы часов — тегами ``{{ hours_* }}``;
стандартные пред-заполненные разделы (ЭБС, ПО, помещения) остаются без изменений.

Скрипт идемпотентный и самодокументирующий: показывает, как именно получены
поставляемые шаблоны. Запуск::

    python -m tools.build_templates --rpd-src "Шаблон РП (2026) ММЭУ.docx" \
        --fos-src "ФОС_ШАБЛОН (2026) ММЭУ.docx"
"""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

from app.core.normalizer import normalize_text


def strip_red_highlights(doc: Document) -> None:
    """Удаляет красные highlight-пометки (в т.ч. на знаках абзаца), унаследованные
    из исходной формы 2026. Жёлтые (наши) не трогаем."""
    for h in list(doc.element.body.iter(qn("w:highlight"))):
        if (h.get(qn("w:val")) or "") == "red":
            h.getparent().remove(h)

PROJECT_ROOT = Path(__file__).resolve().parent.parent
TEMPLATES_DIR = PROJECT_ROOT / "templates"


# --------------------------------------------------------------------------- #
#  Утилиты редактирования docx
# --------------------------------------------------------------------------- #
def set_text(paragraph: Paragraph, text: str) -> None:
    """Заменяет весь текст абзаца одним runом (сохраняя стиль абзаца)."""
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)
    paragraph.add_run(text)


def clear(paragraph: Paragraph) -> None:
    set_text(paragraph, "")


def hl_run(run) -> None:
    """Жёлтая заливка шрифта на run (рендеренное значение тега будет жёлтым)."""
    from docx.enum.text import WD_COLOR_INDEX

    run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def set_value_tag(paragraph: Paragraph, tag: str) -> None:
    """Абзац = только тег значения, жёлтый (рендеренное значение подсветится)."""
    set_text(paragraph, tag)
    hl_run(paragraph.runs[0])


def set_label_value(paragraph: Paragraph, label: str, tag: str) -> None:
    """Абзац = метка (без подсветки) + тег значения (жёлтый)."""
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)
    paragraph.add_run(label)
    hl_run(paragraph.add_run(tag))


def set_label_value_underlined(
    paragraph: Paragraph, label: str, tag: str, *, value_prefix: str = ""
) -> None:
    """Титульная строка формы: метка + подчёркнутое значение (жёлтое) +
    подчёркнутый таб — «нижняя линия» до табстопа, как в исходной форме.

    ``value_prefix`` — пробельный отступ перед значением (в форме это «\xa0\xa0»
    или пробел), тоже подчёркнутый."""
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)
    paragraph.add_run(label)
    if value_prefix:
        pre = paragraph.add_run(value_prefix)
        pre.underline = True
    value = paragraph.add_run(tag)
    value.underline = True
    hl_run(value)
    tail = paragraph.add_run("\t")
    tail.underline = True


def insert_before(paragraph: Paragraph, text: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    np = Paragraph(new_p, paragraph._parent)
    np.add_run(text)
    return np


def remove_paragraph(paragraph: Paragraph) -> None:
    paragraph._element.getparent().remove(paragraph._element)


def insert_after(paragraph: Paragraph, text: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    np = Paragraph(new_p, paragraph._parent)
    np.add_run(text)
    return np


def n(text: str) -> str:
    return normalize_text(text)


# --------------------------------------------------------------------------- #
#  РПД
# --------------------------------------------------------------------------- #
def tag_rpd(src: Path, dst: Path) -> None:
    doc = Document(str(src))

    # --- Абзацы: замена подписей на теги, очистка примеров --- #
    for p in doc.paragraphs:
        t = n(p.text)
        if not t:
            continue
        if "код и наименование дисциплины" in t and p.text.strip().isupper():
            set_value_tag(p, "{{ index }} {{ name }}")
        elif t.startswith("по направлению подготовки"):
            set_label_value(p, "по направлению подготовки ", "{{ direction }}")
        elif t.startswith("направленности"):
            set_label_value(p, "направленности (профиля) ", "{{ profile }}")
        elif t.startswith("форма обучения"):
            set_label_value(p, "форма обучения ", "{{ form_study }}")
        elif "цели в исходной программе нет" in t:
            clear(p)  # целей в исходных РПД нет — раздел заполняется вручную
        # Subdoc-теги — с префиксом «p»: docxtpl убирает абзац-носитель и
        # вставляет блочный XML на уровень body (иначе w:p вкладывается в w:t —
        # невалидный OOXML, контент не виден python-docx и др. потребителям).
        elif "из п. 2 исходной рпд (столбцы 2,3)" in t:
            set_text(p, "{{p competencies }}")
        elif "из п. 2 исходной рпд (столбец 4)" in t:
            set_text(p, "{{p indicators }}")
        elif "из исходной рпд п.3" in t:
            set_text(p, "{{p thematic_plan }}")  # тематический план из старого документа (best-effort)
        elif "указать основную литературу" in t:
            set_text(p, "{{p literature_main }}")
        elif "указать дополнительную литературу" in t:
            set_text(p, "{{p literature_extra }}")
        elif "указать ресурсы сети интернет" in t:
            set_text(p, "{{p internet_resources }}")
        elif "берем из учебного плана" in t:
            clear(p)
        elif t.startswith("заседания кафедры"):
            set_label_value(p, "заседания кафедры ", "{{ department_name }}")
        elif "методов в экономике и управлении" in t:
            clear(p)  # хвост захардкоженного названия кафедры
        elif t.startswith("опк-1"):  # примеры компетенций/индикаторов
            clear(p)

    # --- Таблица часов: значения из Excel --- #
    _tag_hours_table(doc)

    # --- Остальные редакторские пометки (§3 тематический план, §8 оценивание) --- #
    _clean_editorial(doc)
    strip_red_highlights(doc)

    dst.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(dst))
    print(f"[РПД] {src.name} → {dst}")


def _clean_editorial(doc: Document) -> None:
    """Убирает/заполняет оставшиеся редакторские пометки в РПД.

    * §3: убирает инструкции «Указать несколько тем…» и «Проверка итого…»
      (тематический план заполняется автором по официальной пустой таблице);
    * §8: примечание «Исходная РПД п.п.8…» убирается, ячейки «Код/коды
      компетенций» → тег ``{{ competence_codes }}`` (из Базы), инструкции
      «(перечислить)» удаляются.
    """
    # Верхнеуровневые красные примечания.
    for p in doc.paragraphs:
        t = n(p.text)
        if "проверка итого" in t or "исходная рпд п. п. 8" in t or "столбцы 2 и 5" in t:
            clear(p)

    grade_level = {"отлично": "5", "хорошо": "4", "удовлетворительно": "3", "неудовлетворительно": "2"}
    for table in doc.tables:
        header = n(" ".join(c.text for c in table.rows[0].cells))
        if "формируемая компетенция" in header:  # §8 «Система оценивания»
            for row in table.rows:
                level = grade_level.get(n(row.cells[0].text))
                if level is None:
                    continue  # строка-шапка
                # c1 = «Формируемая компетенция» → родительские коды (жёлтые)
                if len(row.cells) > 1:
                    set_value_tag(row.cells[1].paragraphs[0], "{{ competence_parents }}")
                    for extra in row.cells[1].paragraphs[1:]:
                        clear(extra)
                # c2 = «Наименование результатов»: оставить вводную фразу (подсветить),
                # убрать «Знает/Умеет/Владеет (перечислить)», добавить тег результатов.
                if len(row.cells) > 2:
                    cell = row.cells[2]
                    for p in list(cell.paragraphs):
                        if "демонстрирует" in n(p.text):
                            for run in p.runs:
                                if run.text.strip():
                                    hl_run(run)
                        else:  # удаляем «Знает/Умеет/Владеет (перечислить)» без пустых строк
                            p._element.getparent().remove(p._element)
                    hl_run(cell.add_paragraph().add_run("{{ outcomes_" + level + " }}"))
        else:  # §3 «Тематический план» и прочие
            for row in table.rows:
                for cell in row.cells:
                    if "указать несколько тем" in n(cell.text):
                        for p in cell.paragraphs:
                            clear(p)


def _tag_hours_table(doc: Document) -> None:
    # строка-метка (нормализованная) → имя тега для колонки «Всего» (c2) и «семестр» (c3)
    row_tags = {
        "зач. ед.": "{{ ze }}",
        "ак.ч.": "{{ hours_total }}",
        "часы контактной работы": "{{ hours_contact }}",
        "из них аудиторной работы": "{{ hours_aud }}",
        "лекции": "{{ hours_lectures }}",
        "практические занятия": "{{ hours_practical }}",
        "лабораторные занятия": "{{ hours_lab }}",
        "проектное обучение": "{{ hours_project }}",
        "часы внеаудиторной работы": "{{ hours_extra_contact }}",
        "часы самостоятельной работы": "{{ hours_self_study }}",
        "вид промежуточной аттестации": "{{ control_summary }}",
    }
    for table in doc.tables:
        header = " ".join(c.text for c in table.rows[0].cells)
        if "Вид учебной работы" not in header:
            continue
        for row in table.rows:
            # «Общая трудоёмкость» в c0, а единица («зач. ед.»/«ак.ч.») — в c1
            label = n(row.cells[0].text + " " + (row.cells[1].text if len(row.cells) > 1 else ""))
            tag = next((v for k, v in row_tags.items() if k in label), None)
            if tag is None:
                continue
            # колонка «Всего» = индекс 2; «в семестре» = 3 (для односеместровых равны).
            # Значения подсвечиваются жёлтым (всё, что заполнено из Базы).
            if len(row.cells) > 2:
                set_value_tag(row.cells[2].paragraphs[0], tag)
            if len(row.cells) > 3:
                # Для аттестации в семестровой колонке — вид без номера семестра
                # (эталон заполняет обе ячейки: «экзамен | экзамен»).
                sem_tag = "{{ control_kind }}" if "{{ control" in tag else tag
                set_value_tag(row.cells[3].paragraphs[0], sem_tag)
        break


# --------------------------------------------------------------------------- #
#  ФОС
# --------------------------------------------------------------------------- #
def tag_fos(src: Path, dst: Path) -> None:
    doc = Document(str(src))
    gia_heading: Paragraph | None = None
    to_remove: list[Paragraph] = []

    for p in doc.paragraphs:
        t = n(p.text)
        if not t:
            continue
        if t == "код наименование":
            set_value_tag(p, "{{ index }} {{ name }}")
        elif t.startswith("по направлению подготовки"):
            # Пробельные отступы перед значениями — как в форме/эталонах.
            set_label_value_underlined(p, "по направлению подготовки ", "{{ direction }}", value_prefix=" ")
        elif t.startswith("направленности"):
            set_label_value_underlined(p, "направленности (профиля) ", "{{ profile }}")
        elif t.startswith("форма обучения"):
            set_label_value_underlined(p, "форма обучения ", "{{ form_study }}", value_prefix="  ")
        # Статики формы «Примерный перечень задач» / «Примерный состав тестовых
        # вопросов…» сохраняются (так в одобренных эталонах) — контент после них.
        elif "примерный перечень задач" in t:
            pass  # статика формы, контент вставляется после «Задачи к разделу 1…»
        elif "примерный состав тестовых вопросов" in t:
            insert_after(p, "{{p interim_attestation }}")
        elif t.startswith("задачи к разделу 1"):
            # Первая строка формы остаётся: индикаторы — из Базы (жёлтым),
            # разбивку по разделам уточняет методист. Контент задач — после неё.
            set_label_value(p, "Задачи к разделу 1. (оцениваемая компетенция и индикатор ", "{{ fos_indicators }}")
            p.add_run(")")
            insert_after(p, "{{p current_control }}")
        elif t.startswith("задачи к разделу"):
            to_remove.append(p)  # вторая строка-пример формы
        elif t.startswith("задача") or t == "ответ:":
            to_remove.append(p)  # примеры формы: удаляем абзац целиком, не оставляя пустот
        elif "обязательно с ответами" in t:
            to_remove.append(p)
        elif t.startswith("заседания кафедры"):
            set_label_value_underlined(p, "заседания кафедры ", "{{ department_name }}", value_prefix="  ")
        elif "методов в экономике и управлении" in t:
            to_remove.append(p)  # хвост захардкоженного названия кафедры (вторая строка)
        elif "государственной итоговой" in t:
            gia_heading = p

    if gia_heading is not None:
        # Раздела 3 нет вовсе, если дисциплина не участвует в ГИА (исходный ФОС
        # без блока ГИА): заголовок + контент под jinja-условием.
        insert_before(gia_heading, "{%p if gia_present %}")
        tag_p = insert_after(gia_heading, "{{p gia }}")
        insert_after(tag_p, "{%p endif %}")
    for p in to_remove:
        remove_paragraph(p)
    _tighten_fos_title(doc)

    strip_red_highlights(doc)
    dst.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(dst))
    print(f"[ФОС] {src.name} → {dst}")


def _tighten_fos_title(doc: Document) -> None:
    """Убирает 2 пустых абзаца перед «Москва, …» на титуле ФОС.

    Реальные значения (название дисциплины, профиль) длиннее однострочных
    заглушек формы и переносятся — без запаса «Москва, 2026» уезжает на
    следующую страницу."""
    paragraphs = doc.paragraphs
    for i, p in enumerate(paragraphs):
        if n(p.text).startswith("москва"):
            removed = 0
            for q in reversed(paragraphs[:i]):
                if removed == 2:
                    break
                if not q.text.strip():
                    remove_paragraph(q)
                    removed += 1
                else:
                    break
            break


def main() -> None:
    ap = argparse.ArgumentParser(description="Разметка шаблонов 2026 тегами docxtpl")
    _src_dir = PROJECT_ROOT / "files" if (PROJECT_ROOT / "files").is_dir() else PROJECT_ROOT
    ap.add_argument("--rpd-src", default=str(_src_dir / "Шаблон РП (2026) ММЭУ.docx"))
    ap.add_argument("--fos-src", default=str(_src_dir / "ФОС_ШАБЛОН (2026) ММЭУ.docx"))
    args = ap.parse_args()

    tag_rpd(Path(args.rpd_src), TEMPLATES_DIR / "rpd_2026_tagged.docx")
    tag_fos(Path(args.fos_src), TEMPLATES_DIR / "fos_2026_tagged.docx")


if __name__ == "__main__":
    main()
