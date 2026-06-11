"""Асинхронный ETL-оркестратор EduConvert (ТЗ §4, §7).

Связывает слои в конвейер пакетной обработки:

#. инициализация — загрузка Базы и валидация шаблонов;
#. для каждого файла (с **изоляцией сбоев**): индекс → поиск в Базе →
   конвертация .doc → diff старых чисел с эталоном → извлечение контента →
   генерация документа;
#. формирование ``report.xlsx`` и упаковка результатов в ZIP;
#. очистка временных файлов после выдачи архива.

Все блокирующие операции (чтение/запись docx, конвертация .doc, чтение Excel,
архивирование) выполняются в пуле потоков через :func:`anyio.to_thread.run_sync`,
чтобы не блокировать event loop (ТЗ §2, §7.1). Ошибка одного файла не прерывает
обработку остальных.
"""

from __future__ import annotations

import logging
import shutil
import uuid
import zipfile
from collections.abc import Callable
from dataclasses import dataclass, field
from pathlib import Path

import anyio
import polars as pl

from app.config import settings
from app.core.differ import compute_diffs, hours_consistency_check
from app.core.doc_converter import WordComConverter
from app.core.excel_parser import ExcelSubjectRepository, load_repository
from app.core.exceptions import (
    EduConvertError,
    SubjectNotFoundError,
    TemplateValidationError,
)
from app.core.models import (
    DocType,
    FileResult,
    FileStatus,
    RunReport,
)
from app.core.normalizer import index_from_filename, normalize_text
from app.core.word_extractor import WordExtractor
from app.core.word_generator import DocxtplGenerator

logger = logging.getLogger("educonvert.orchestrator")

ProgressCallback = Callable[[int, int, str], None]

# Колонки сводного отчёта — дословно по ТЗ §8.
_REPORT_COLUMNS = [
    "Имя исходного файла",
    "Индекс дисциплины",
    "Статус",
    "Описание ошибки или расхождения",
    "Значение в старом файле",
    "Значение в Базе",
]


@dataclass
class RunResult:
    """Итог прогона: отчёт, путь к ZIP и временная папка для очистки."""

    report: RunReport
    zip_path: Path
    workdir: Path

    def cleanup(self) -> None:
        """Удаляет временную папку и архив (ТЗ §7.3)."""
        shutil.rmtree(self.workdir, ignore_errors=True)
        self.zip_path.unlink(missing_ok=True)


class Orchestrator:
    """Оркестратор пакетной конвертации РПД/ФОС."""

    def __init__(
        self,
        db_path: str | Path,
        rpd_template: str | Path | None = None,
        fos_template: str | Path | None = None,
        *,
        repository: ExcelSubjectRepository | None = None,
        converter: WordComConverter | None = None,
        extractor: WordExtractor | None = None,
        generator: DocxtplGenerator | None = None,
    ) -> None:
        self.db_path = Path(db_path)
        self.rpd_template = Path(rpd_template or settings.rpd_template)
        self.fos_template = Path(fos_template or settings.fos_template)
        self._repo = repository
        self.converter = converter or WordComConverter()
        self.extractor = extractor or WordExtractor()
        self.generator = generator or DocxtplGenerator()

    # ------------------------------------------------------------------ #
    #  Публичный запуск
    # ------------------------------------------------------------------ #
    async def run(
        self,
        input_files: list[Path],
        progress: ProgressCallback | None = None,
    ) -> RunResult:
        """Обрабатывает список файлов и возвращает отчёт + ZIP-архив."""

        def notify(done: int, total: int, message: str) -> None:
            if progress is not None:
                progress(done, total, message)

        workdir = settings.temp_root / f"run_{uuid.uuid4().hex[:12]}"
        out_dir = workdir / "output"
        out_dir.mkdir(parents=True, exist_ok=True)

        try:
            notify(0, len(input_files), "Инициализация: загрузка Базы и проверка шаблонов…")
            await anyio.to_thread.run_sync(self._initialize, workdir)
            report = RunReport()
            total = len(input_files)
            for i, path in enumerate(input_files):
                notify(i, total, f"Обработка файла: {path.name}")
                result = await anyio.to_thread.run_sync(self._process_file_safe, path, out_dir)
                report.results.append(result)

            notify(total, total, "Формирование отчёта и архива…")
            report_path = await anyio.to_thread.run_sync(self._write_report, report, out_dir)
            zip_path = await anyio.to_thread.run_sync(self._make_zip, out_dir, workdir)
        except BaseException:
            # Сбой до выдачи RunResult — вызвать cleanup() будет некому,
            # временную папку убираем сами (ТЗ §7.3).
            shutil.rmtree(workdir, ignore_errors=True)
            raise
        logger.info(
            "Готово: успешно=%d, расхождений=%d, ошибок=%d",
            report.succeeded,
            report.with_discrepancies,
            report.failed,
        )
        _ = report_path
        return RunResult(report=report, zip_path=zip_path, workdir=workdir)

    # ------------------------------------------------------------------ #
    #  Инициализация
    # ------------------------------------------------------------------ #
    def _initialize(self, workdir: Path | None = None) -> None:
        if self._repo is None:
            self._repo = load_repository(self.db_path, settings.plan_sheet)
        self.rpd_template = self._ensure_tagged(self.rpd_template, DocType.RPD, workdir)
        self.fos_template = self._ensure_tagged(self.fos_template, DocType.FOS, workdir)

    def _ensure_tagged(self, path: Path, doc_type: DocType, workdir: Path | None) -> Path:
        """Возвращает размеченный шаблон; чистую официальную форму размечает на лету.

        Пользователь может выбрать в качестве шаблона саму форму 2026 (без
        docxtpl-тегов) — тогда она автоматически размечается во временную копию
        тем же кодом, что собирает поставляемые ``templates/*_tagged.docx``.
        """
        try:
            self.generator.validate_template(path, doc_type)
            return path
        except TemplateValidationError as exc:
            if workdir is None or not _looks_like_official_form(path, doc_type):
                raise EduConvertError(
                    f"В шаблоне '{path.name}' нет docxtpl-тегов (не хватает: "
                    f"{', '.join(exc.missing_tags)}), и на официальную форму 2026 "
                    f"он не похож. Выберите размеченный шаблон из templates/ "
                    f"({'rpd' if doc_type == DocType.RPD else 'fos'}_2026_tagged.docx) "
                    f"или саму форму 2026 — её конвертер разметит автоматически."
                ) from exc
            from tools.build_templates import tag_fos, tag_rpd  # noqa: PLC0415 — ленивый импорт

            tagged = workdir / f"autotag_{doc_type.value}_{path.stem}.docx"
            (tag_rpd if doc_type == DocType.RPD else tag_fos)(path, tagged)
            self.generator.validate_template(tagged, doc_type)
            logger.info(
                "Шаблон '%s' без тегов — выполнена авторазметка официальной формы 2026",
                path.name,
            )
            return tagged

    @property
    def repository(self) -> ExcelSubjectRepository:
        if self._repo is None:
            self._initialize()
        assert self._repo is not None
        return self._repo

    # ------------------------------------------------------------------ #
    #  Обработка одного файла (изоляция сбоев)
    # ------------------------------------------------------------------ #
    def _process_file_safe(self, path: Path, out_dir: Path) -> FileResult:
        try:
            return self._process_file(path, out_dir)
        except SubjectNotFoundError as exc:
            logger.warning("Файл %s: %s", path.name, exc)
            return FileResult(
                filename=path.name,
                index=exc.index,
                doc_type=_detect_doc_type(path),
                status=FileStatus.ERROR,
                message=str(exc),
            )
        except EduConvertError as exc:
            logger.error("Файл %s: %s", path.name, exc)
            return FileResult(
                filename=path.name,
                doc_type=_detect_doc_type(path),
                status=FileStatus.ERROR,
                message=str(exc),
            )
        except Exception as exc:  # noqa: BLE001 — изоляция: любой сбой → строка ошибки
            logger.exception("Непредвиденная ошибка при обработке %s", path.name)
            return FileResult(
                filename=path.name,
                doc_type=_detect_doc_type(path),
                status=FileStatus.ERROR,
                message=f"Непредвиденная ошибка: {exc}",
            )

    def _process_file(self, path: Path, out_dir: Path) -> FileResult:
        doc_type = _detect_doc_type(path)

        # 1. Индекс — сначала из имени файла (не требует открытия документа).
        index = index_from_filename(path.name)

        # 2. Конвертация устаревшего .doc.
        docx_path = self.converter.convert(path, out_dir / "_converted") if self.converter.supports(
            path
        ) else path

        if index is None:
            index = self.extractor.extract_index(docx_path)

        # 3. Поиск в Базе (источник истины).
        subject = self.repository.get_subject(index) if index else None
        if subject is None:
            raise SubjectNotFoundError(index or path.name)

        # 4. Контент из старого документа.
        content = self.extractor.extract(docx_path, doc_type)

        # 5. Diff старых чисел с эталоном (только для РПД — там есть таблица часов)
        # + проверка целостности часов самой Базы (для всех типов).
        diffs = []
        if doc_type == DocType.RPD:
            old = self.extractor.extract_old_numbers(docx_path)
            diffs = compute_diffs(old, subject)
        hours_check = hours_consistency_check(subject)
        if hours_check is not None:
            diffs.append(hours_check)

        # 6. Генерация (числа из Excel, текст из Word).
        template = self.rpd_template if doc_type == DocType.RPD else self.fos_template
        out_name = f"{subject.index}_{doc_type.value}_2026.docx"
        self.generator.generate(template, out_dir / out_name, subject, content, doc_type)

        status = FileStatus.DISCREPANCY if diffs else FileStatus.SUCCESS
        return FileResult(
            filename=path.name,
            index=subject.index,
            doc_type=doc_type,
            status=status,
            message="Расхождения чисел (в документ записаны значения из Базы)"
            if diffs
            else "",
            diffs=diffs,
            output_name=out_name,
        )

    # ------------------------------------------------------------------ #
    #  Отчёт и упаковка
    # ------------------------------------------------------------------ #
    def _write_report(self, report: RunReport, out_dir: Path) -> Path:
        rows: list[dict[str, str]] = []
        for r in report.results:
            base = {"Имя исходного файла": r.filename, "Индекс дисциплины": r.index or ""}
            desc = "Описание ошибки или расхождения"
            if r.status == FileStatus.ERROR:
                rows.append({**base, "Статус": r.status.value, desc: r.message,
                             "Значение в старом файле": "", "Значение в Базе": ""})
            elif r.diffs:
                for d in r.diffs:
                    rows.append({**base, "Статус": r.status.value, desc: d.field,
                                 "Значение в старом файле": d.old_value,
                                 "Значение в Базе": d.new_value})
            else:
                rows.append({**base, "Статус": r.status.value, desc: "",
                             "Значение в старом файле": "", "Значение в Базе": ""})

        if not rows:
            rows.append(dict.fromkeys(_REPORT_COLUMNS, ""))
        df = pl.DataFrame(rows).select(_REPORT_COLUMNS)
        report_path = out_dir / "report.xlsx"
        df.write_excel(report_path, autofit=True)
        return report_path

    def _make_zip(self, out_dir: Path, workdir: Path) -> Path:
        zip_path = workdir.parent / f"{workdir.name}.zip"
        with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
            for file in sorted(out_dir.rglob("*")):
                if file.is_file() and "_converted" not in file.parts:
                    zf.write(file, file.relative_to(out_dir))
        return zip_path


def _looks_like_official_form(path: Path, doc_type: DocType) -> bool:
    """Похож ли файл на чистую официальную форму 2026 (кандидат на авторазметку)."""
    try:
        from docx import Document  # noqa: PLC0415 — только для этой проверки

        doc = Document(str(path))
    except Exception:  # noqa: BLE001 — не docx/битый файл — не форма
        return False
    text = normalize_text(" ".join(p.text for p in doc.paragraphs))
    if doc_type == DocType.RPD:
        tables = normalize_text(
            " ".join(c.text for t in doc.tables for r in t.rows[:2] for c in r.cells)
        )
        return "цели освоения дисциплины" in text and "вид учебной работы" in tables
    return "примерный перечень задач" in text or "примерный состав тестовых вопросов" in text


def _detect_doc_type(path: Path) -> DocType:
    name = normalize_text(path.name)
    if name.startswith("фос") or "фос_" in name or "_фос_" in name:
        return DocType.FOS
    if "рпд" in name or "рп_" in name:
        return DocType.RPD
    return DocType.RPD
