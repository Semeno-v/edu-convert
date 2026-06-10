"""Генерация целевого документа из шаблона docxtpl (ТЗ §4 Этап 2.6).

Соответствует одобренному эталонному формату:

* **Всё заполненное конвертацией выделяется жёлтым** (заливка шрифта). Статичные
  разделы шаблона (заголовки, §5/§6/§7) остаются без заливки.
* **§2** — компетенции и индикаторы строятся из Базы (лист «Компетенции»,
  :attr:`SubjectData.competencies`) как текстовые абзацы.
* **§3** — числа из Базы (подсветка задаётся в шаблоне на тегах), тематический
  план — из старого документа (best-effort).
* **§4** — литература переформатируется из таблицы старого документа в
  нумерованный список «Автор. Название. Выходные данные. URL».
* **§8** — «Формируемая компетенция» = родительские коды (ПК-1, ПК-2).

Подсветка скалярных значений (титул, часы, коды) задаётся в самом шаблоне
(жёлтый highlight на run-тегах), подсветка subdoc-контента — здесь.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.shared import Pt
from docxtpl import DocxTemplate

from app.config import settings
from app.core.exceptions import TemplateValidationError
from app.core.models import (
    ContentBlock,
    ContentBlocks,
    DocType,
    ElementKind,
    RichParagraph,
    RichTable,
    SubjectData,
)
from app.core.normalizer import normalize_text

_PLACEHOLDER = "Не предусмотрено."


class DocxtplGenerator:
    """Рендерер целевых документов на базе docxtpl."""

    def validate_template(self, template_path: Path, doc_type: DocType) -> None:
        tpl = DocxTemplate(str(template_path))
        try:
            found = tpl.get_undeclared_template_variables()
        except Exception:  # noqa: BLE001 — fallback на текстовый поиск тегов
            found = self._scan_tags(template_path)
        required = (
            settings.required_rpd_tags
            if doc_type == DocType.RPD
            else settings.required_fos_tags
        )
        missing = [tag for tag in required if tag not in found]
        if missing:
            raise TemplateValidationError(missing, template_path.name)

    @staticmethod
    def _scan_tags(template_path: Path) -> set[str]:
        import zipfile

        with zipfile.ZipFile(template_path) as zf:
            xml = zf.read("word/document.xml").decode("utf-8", "replace")
        text = re.sub(r"<[^>]+>", "", xml)
        # Учитываем префиксы docxtpl: {{r …}} (RichText) и {{p …}} (блочный subdoc).
        return set(re.findall(r"\{\{\s*(?:[pr]\s+)?([a-zA-Z_][a-zA-Z0-9_]*)", text))

    def generate(
        self,
        template_path: Path,
        out_path: Path,
        subject: SubjectData,
        content: ContentBlocks,
        doc_type: DocType,
    ) -> Path:
        tpl = DocxTemplate(str(template_path))
        context = self._build_context(tpl, subject, content, doc_type)
        tpl.render(context, autoescape=True)
        if doc_type == DocType.RPD:
            # Официальная таблица тем заполняется после рендера: темы — из
            # исходника, часы масштабируются к Базе (золотое правило).
            # Именно tpl.docx: get_docx() перезагрузил бы исходный шаблон,
            # затерев отрендеренное дерево (init_docx при is_rendered).
            _fill_topics_table(tpl.docx, content.get("thematic_plan"), subject)
            _apply_body_justify(tpl.docx)
        out_path.parent.mkdir(parents=True, exist_ok=True)
        tpl.save(str(out_path))
        return out_path

    # ------------------------------------------------------------------ #
    #  Сборка контекста
    # ------------------------------------------------------------------ #
    def _build_context(
        self,
        tpl: DocxTemplate,
        subject: SubjectData,
        content: ContentBlocks,
        doc_type: DocType,
    ) -> dict[str, object]:
        # Скалярные значения; их жёлтая подсветка задаётся в шаблоне на тегах.
        context: dict[str, object] = {
            "index": subject.index,
            "name": subject.name,
            "ze": _num(subject.ze),
            "hours_total": subject.hours_total,
            "hours_contact": subject.hours_contact,
            "hours_aud": subject.hours_aud,
            # Нулевые виды занятий в таблице часов — прочерк (как в эталонах).
            "hours_lectures": subject.hours_lectures or "-",
            "hours_practical": subject.hours_practical or "-",
            "hours_lab": subject.hours_lab or "-",
            "hours_project": subject.hours_project or "-",
            "hours_extra_contact": max(subject.hours_contact - subject.hours_aud, 0),
            # «Часы самостоятельной работы» в форме = чистая СРС из Базы
            # (часы контроля в таблице эталона не отображаются).
            "hours_self_study": subject.hours_srs,
            "control_summary": subject.control_summary,
            # Вид аттестации строчными, без семестра — обе ячейки таблицы §3
            # эталона: «экзамен | экзамен».
            "control_kind": ", ".join(dict.fromkeys(cf.kind.value for cf in subject.control_forms)).lower(),
            "semesters": ", ".join(str(s) for s in subject.semesters),
            "department": subject.department or "",
            # «заседания кафедры математических методов…» — род. падеж контекста,
            # первая буква строчная (как в эталонах); аббревиатуры не трогаем.
            "department_name": _decapitalize(subject.department_name or ""),
            # §8 — родительские коды компетенций (ПК-1, ПК-2).
            "competence_parents": ", ".join(subject.competence_parents),
            # Титул: направление/профиль/форма — атрибуты программы из листа
            # «Титул» Базы (титул старого документа часто скопирован с чужой
            # программы); старый документ — только фолбэк.
            "direction": subject.direction or content.direction or "",
            "profile": subject.profile or content.profile or "",
            "form_study": subject.form_study or content.form_study or "очная",
        }

        if doc_type == DocType.RPD:
            # §8: результаты обучения по уровням оценки — из §8 исходной РПД.
            outcomes = _assessment_outcomes(content.get("assessment"))
            for level in ("5", "4", "3", "2"):
                context["outcomes_" + level] = outcomes.get(level, "")
            context["competencies"] = self._competencies_subdoc(tpl, subject)
            context["indicators"] = self._indicators_subdoc(tpl, subject)
            context["literature_main"] = self._literature_subdoc(tpl, content.get("literature_main"))
            context["literature_extra"] = self._literature_subdoc(tpl, content.get("literature_extra"))
            # §4.3: только настоящие интернет-ресурсы (MS Teams и т.п.) —
            # ссылки на ЭБС дублируют статичный §5 и отфильтровываются
            # (эталоны: «Не предусмотрено.» либо нумерованный список).
            context["internet_resources"] = self._internet_resources_subdoc(
                tpl, content.get("internet_resources")
            )
        else:
            # Индикаторы для строки «Задачи к разделу 1. (…индикатор …)» — из Базы.
            context["fos_indicators"] = ", ".join(subject.competence_codes)
            for key in ("current_control", "interim_attestation", "gia"):
                # Контент ФОС в эталонах переформатирован: по ширине + отступ 1 см.
                context[key] = self._block_to_subdoc(tpl, content.get(key), body_format=True)
            # Раздел 3 (ГИА) выводится только при наличии контента в исходнике:
            # составить выборку задач для ГИА конвертер за методиста не может.
            gia = content.get("gia")
            context["gia_present"] = gia is not None and not gia.is_empty
        return context

    # ------------------------------------------------------------------ #
    #  §2 — компетенции и индикаторы из Базы (текстовые абзацы, жёлтые)
    # ------------------------------------------------------------------ #
    def _competencies_subdoc(self, tpl: DocxTemplate, subject: SubjectData):
        sd = tpl.new_subdoc()
        if not subject.competencies:
            self._add_text(sd, _PLACEHOLDER, highlight=True, font_size_pt=12)
            return sd
        for group in subject.competencies:
            self._add_text(sd, _dot(f"{group.code} - {group.text}".strip(" -")), highlight=True, font_size_pt=12)
        return sd

    def _indicators_subdoc(self, tpl: DocxTemplate, subject: SubjectData):
        sd = tpl.new_subdoc()
        indicators = [ind for g in subject.competencies for ind in g.indicators]
        if not indicators:
            self._add_text(sd, _PLACEHOLDER, highlight=True, font_size_pt=12)
            return sd
        for ind in indicators:
            self._add_text(sd, _dot(f"{ind.code}. {ind.text}".strip()), highlight=True, font_size_pt=12)
        return sd

    # ------------------------------------------------------------------ #
    #  §4 — литература: таблица старого документа → нумерованный список
    # ------------------------------------------------------------------ #
    def _literature_subdoc(self, tpl: DocxTemplate, block: ContentBlock | None):
        sd = tpl.new_subdoc()
        citations = _table_to_citations(block)
        if not citations:
            self._add_text(sd, _PLACEHOLDER, highlight=True, font_size_pt=12)
            return sd
        for i, cite in enumerate(citations, 1):
            self._add_text(sd, f"{i}. {cite}", highlight=True, font_size_pt=12)
        return sd

    # ЭБС, уже перечисленные в статичном §5 шаблона: их адреса в §4.3 — дубли.
    _EBS_MARKERS = (
        "znanium", "biblioclub", "urait", "юрайт", "book.ru", "lanbook",
        "grebennikon", "ibooks", "мтс линк", "линк курсы",
    )

    def _internet_resources_subdoc(self, tpl: DocxTemplate, block: ContentBlock | None):
        """§4.3: нумерованный список настоящих интернет-ресурсов исходника."""
        sd = tpl.new_subdoc()
        items: list[str] = []
        for element in (block.elements if block else []):
            if element.kind == ElementKind.TABLE and element.table is not None:
                for row in element.table.rows:
                    cells = [_cell_text(c).strip() for c in row.cells]
                    texts = [c for c in cells if c and c not in ("-",)]
                    if not texts:
                        continue
                    joined_l = " ".join(texts).lower()
                    if "наименование" in joined_l and ("адрес" in joined_l or "доступ" in joined_l):
                        continue  # шапка таблицы
                    if any(m in joined_l for m in self._EBS_MARKERS):
                        continue  # дубль ЭБС из §5
                    if not any(ch.isalpha() for ch in joined_l):
                        continue
                    body = [c for c in texts if not re.fullmatch(r"\d+\.?", c)]  # «№ п/п»
                    cite = re.sub(r"\.\s*\.", ".", ". ".join(body))
                    items.append(cite if cite.endswith(".") else cite + ".")
            elif element.kind == ElementKind.PARAGRAPH and element.paragraph is not None:
                text = element.paragraph.text.strip()
                if "http" in text.lower() and not any(m in text.lower() for m in self._EBS_MARKERS):
                    items.append(text)
        if not items:
            self._add_text(sd, _PLACEHOLDER, highlight=True, font_size_pt=12)
            return sd
        for i, item in enumerate(items, 1):
            self._add_text(sd, f"{i}. {item}", highlight=True, font_size_pt=12)
        return sd

    # ------------------------------------------------------------------ #
    #  Общие subdoc-блоки (контент из старого документа, жёлтый)
    # ------------------------------------------------------------------ #
    def _block_to_subdoc(
        self,
        tpl: DocxTemplate,
        block: ContentBlock | None,
        empty_text: str = _PLACEHOLDER,
        *,
        body_format: bool = False,
        font_size_pt: float | None = None,
    ):
        sd = tpl.new_subdoc()
        if block is None or block.is_empty:
            self._add_text(sd, empty_text, highlight=bool(empty_text), font_size_pt=font_size_pt)
            return sd
        for element in block.elements:
            if element.kind == ElementKind.PARAGRAPH and element.paragraph is not None:
                self._add_paragraph(
                    sd, element.paragraph, highlight=True, body_format=body_format,
                    font_size_pt=font_size_pt,
                )
            elif element.kind == ElementKind.TABLE and element.table is not None:
                self._add_table(sd, element.table, highlight=True)
        return sd

    # ------------------------------------------------------------------ #
    #  Низкоуровневая вставка с подсветкой
    # ------------------------------------------------------------------ #
    @staticmethod
    def _hl(run, highlight: bool) -> None:
        if highlight:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    def _add_text(self, sd, text: str, *, highlight: bool, font_size_pt: float | None = None) -> None:
        run = sd.add_paragraph().add_run(text)
        if font_size_pt is not None:
            run.font.size = Pt(font_size_pt)
        self._hl(run, highlight)

    _ALIGNMENTS = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }

    # Формат тела ФОС в одобренных эталонах: по ширине, красная строка 1 см.
    _BODY_FIRST_LINE_PT = 28.35

    def _add_paragraph(
        self, sd, rich: RichParagraph, *, highlight: bool, body_format: bool = False,
        font_size_pt: float | None = None,
    ) -> None:
        # Маркер/номер списка уже синтезирован экстрактором в текст ранов —
        # стиль списка не назначаем, чтобы не получить двойной маркер.
        paragraph = sd.add_paragraph()
        if body_format and rich.alignment != "center":
            # Эталонный формат тела (центрированные подзаголовки сохраняются):
            # по ширине, красная строка 1 см, без интервала после абзаца —
            # иначе Normal шаблона добавляет 8pt на каждый абзац (+страница).
            paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph.paragraph_format.first_line_indent = Pt(self._BODY_FIRST_LINE_PT)
            paragraph.paragraph_format.space_after = Pt(0)
        else:
            # Прямое форматирование исходника (выравнивание/отступы) переносится;
            # None — наследование стиля шаблона.
            if rich.alignment in self._ALIGNMENTS:
                paragraph.paragraph_format.alignment = self._ALIGNMENTS[rich.alignment]
            if rich.first_line_indent_pt is not None:
                paragraph.paragraph_format.first_line_indent = Pt(rich.first_line_indent_pt)
            if rich.left_indent_pt is not None:
                paragraph.paragraph_format.left_indent = Pt(rich.left_indent_pt)
        for run in rich.runs:
            r = paragraph.add_run(run.text)
            if not body_format:
                # body_format (ФОС): инлайн-форматирование исходника снимается —
                # эталоны приводят контент к единому стилю формы (жирные
                # «Ответ:» и пр. в одобренных документах вычищены).
                r.bold, r.italic, r.underline = run.bold, run.italic, run.underline
            if font_size_pt is not None:
                r.font.size = Pt(font_size_pt)
            self._hl(r, highlight)

    def _add_table(self, sd, rich: RichTable, *, highlight: bool) -> None:
        if not rich.rows:
            return
        ncols = max((len(row.cells) for row in rich.rows), default=0)
        if ncols == 0:
            return
        table = sd.add_table(rows=len(rich.rows), cols=ncols)
        try:
            table.style = "Table Grid"
        except KeyError:
            pass
        # Сначала восстанавливаем объединения (как в исходнике), потом заполняем.
        for ri, row in enumerate(rich.rows):
            for ci, cell in enumerate(row.cells):
                if cell.merged or (cell.colspan == 1 and cell.rowspan == 1) or ci >= ncols:
                    continue
                br = min(ri + cell.rowspan - 1, len(rich.rows) - 1)
                bc = min(ci + cell.colspan - 1, ncols - 1)
                try:
                    table.cell(ri, ci).merge(table.cell(br, bc))
                except Exception:  # noqa: BLE001 — кривое объединение не валит перенос
                    pass
        # Заполняются только ячейки-истоки: накрытые объединением позиции
        # помечены merged и пропускаются (table.cell для них вернул бы исток).
        # id()-кэш здесь недопустим: lxml-прокси временные, их адреса
        # переиспользуются — ячейки ложно считались уже заполненными.
        for ri, row in enumerate(rich.rows):
            for ci, cell in enumerate(row.cells):
                if cell.merged or ci >= ncols:
                    continue
                tc = table.cell(ri, ci)
                tc.text = ""
                for pi, para in enumerate(cell.paragraphs):
                    p = tc.paragraphs[0] if pi == 0 else tc.add_paragraph()
                    for run in para.runs:
                        r = p.add_run(run.text)
                        r.bold, r.italic = run.bold, run.italic
                        self._hl(r, highlight)


def _cell_text(cell) -> str:
    return " ".join(p.text for p in cell.paragraphs).strip()


def _assessment_outcomes(block: ContentBlock | None) -> dict[str, str]:
    """Результаты обучения для §8 из таблицы оценивания исходной РПД.

    Формат одобренных эталонов: формулировки **лучшего** уровня («Отлично» /
    «Зачтено») агрегируются в абзацы «Знает …; …», «Умеет …; …», «Владеет …; …»
    (повторный глагол у продолжений срезается) и тиражируются на все четыре
    строки §8 — различаются только вводные фразы шаблона. Ячейка уровня в
    исходнике обычно объединена по вертикали — пустой уровень наследуется.
    """
    if block is None:
        return {}
    table = next((e.table for e in block.elements if e.table is not None), None)
    if table is None or len(table.rows) < 2:
        return {}

    grades: dict[str, list[str]] = {}
    last_grade = ""
    for row in table.rows[1:]:
        if not row.cells:
            continue
        grade = normalize_text(_cell_text(row.cells[0])) or last_grade
        last_grade = grade
        requirement = _cell_text(row.cells[-1])
        if grade and requirement and requirement != grade:
            grades.setdefault(grade, []).append(requirement)

    # Лучший уровень: «отлично» / «зачтено» / «зачет» — но не «не зачтено».
    best = next(
        (grades[g] for key in ("отл", "зач") for g in grades
         if key in g and not g.startswith("не")),
        None,
    )
    if not best:
        return {}

    # Группировка по ведущему глаголу; продолжения — через «;» без глагола.
    order = ("знает", "умеет", "владеет")
    groups: dict[str, list[str]] = {k: [] for k in order}
    extras: list[str] = []
    for req in dict.fromkeys(best):
        req = req.strip().rstrip(".")  # точки на стыках «;» эталон срезает
        verb = next((v for v in order if normalize_text(req).startswith(v)), None)
        if verb is None:
            extras.append(req)
        elif groups[verb]:
            groups[verb].append(re.sub(r"^\s*\S+\s+", "", req))  # срезаем глагол
        else:
            groups[verb].append(req)
    paragraphs = ["; ".join(items) for items in groups.values() if items] + extras
    text = "\a".join(paragraphs)  # \a — новый абзац в docxtpl
    return {"5": text, "4": text, "3": text, "2": text}


def _table_to_citations(block: ContentBlock | None) -> list[str]:
    """Переформатирует таблицу литературы в список «Автор. Название. Изд. URL»."""
    if block is None:
        return []
    table = next((e.table for e in block.elements if e.table is not None), None)
    if table is None or len(table.rows) < 2:
        return []
    header = [_cell_text(c).lower() for c in table.rows[0].cells]

    def col(*keywords: str) -> int | None:
        return next((i for i, h in enumerate(header) if any(k in h for k in keywords)), None)

    ca, ct, co, cu = col("автор"), col("наименован", "назван"), col("выходн", "издат", "год"), col("url", "адрес", "ссылк", "доступ")
    citations: list[str] = []
    for row in table.rows[1:]:
        cells = [_cell_text(c) for c in row.cells]

        def g(i: int | None, _cells: list[str] = cells) -> str:
            return _cells[i].strip() if i is not None and i < len(_cells) else ""

        title = re.sub(r"\s+-\s+(\d+-е изд)", r". \1", g(ct))  # «… - 1-е изд» → «…. 1-е изд»
        # Выходные данные — в формат эталона: «2015. 482 с.»
        imprint = g(co)
        imprint = re.sub(r"Год издания:\s*(\d{4})\.?", r"\1.", imprint)
        imprint = re.sub(r"(?:Объем|Кол-во страниц):\s*(\d+)\s*(?:стр\.?|с\.?)?", r"\1 с.", imprint)
        url = g(cu)
        if url and "http" in url and not url.lower().startswith("url"):
            url = "URL: " + url

        parts = [p for p in (g(ca), title, imprint, url) if p]
        if g(ca) or title:
            cite = re.sub(r"\.\s*\.", ".", ". ".join(parts))  # убираем двойные точки
            citations.append(cite)
    return citations


def _to_num(text: str) -> float | None:
    """«1,6» → 1.6; пустые/нечисловые → None."""
    t = text.strip().replace("\xa0", "").replace(",", ".")
    if not t or t == "-":
        return None
    try:
        return float(t)
    except ValueError:
        return None


def _fmt_hours(value: float) -> str:
    """1.6 → «1,6», 4.0 → «4» (формат часов в эталонной таблице тем)."""
    value = round(value, 2)
    if float(value).is_integer():
        return str(int(value))
    return f"{value:g}".replace(".", ",")


def _fill_topics_table(docx, block: ContentBlock | None, subject: SubjectData) -> None:
    """Заполняет официальную таблицу «Темы (разделы) дисциплины» (§3).

    Темы и распределение часов берутся из таблицы §3 исходника, но часы
    каждого вида масштабируются так, чтобы итоги совпали с Базой (в старых
    таблицах часы устаревшие — золотое правило). Колонки исходника ищутся по
    ключевым словам шапки; строки-болванки шаблона заменяются данными.
    """
    table = next(
        (
            t
            for t in docx.tables
            if len(t.columns) == 7 and "Темы (разделы)" in t.rows[0].cells[1].text
        ),
        None,
    )
    src = next((e.table for e in block.elements if e.table is not None), None) if block else None
    if table is None or src is None or len(src.rows) < 2:
        return

    def cell_text(cell) -> str:
        # Перенос строки внутри ячейки исходника (несколько абзацев) сохраняется.
        return "\n".join(p.text for p in cell.paragraphs).strip()

    # Колонки исходной таблицы по ключевым словам шапки (первые 3 строки).
    nsrc = max(len(r.cells) for r in src.rows)
    labels = ["" for _ in range(nsrc)]
    for row in src.rows[:3]:
        for c, cell in enumerate(row.cells[:nsrc]):
            labels[c] += " " + normalize_text(cell_text(cell))

    def find_col(*keywords: str) -> int | None:
        return next((c for c, lab in enumerate(labels) if any(k in lab for k in keywords)), None)

    col_topic = find_col("содержан", "наименование тем", "темы")
    cols = {
        "lectures": (find_col("лекци"), subject.hours_lectures),
        "practical": (find_col("практическ", "семинар"), subject.hours_practical),
        "lab": (find_col("лабораторн"), subject.hours_lab),
        "project": (find_col("проектн"), subject.hours_project),
    }
    if col_topic is None or cols["lectures"][0] is None:
        return

    # Строки тем: непустая тема, не шапка и не «итого».
    topic_rows = []
    for row in src.rows:
        cells = [cell_text(c) for c in row.cells]
        topic = cells[col_topic] if col_topic < len(cells) else ""
        if not topic or "итого" in normalize_text(cells[0]) or "итого" in normalize_text(topic):
            continue
        if normalize_text(topic) in ("содержание", "наименование тем (разделов)"):
            continue
        topic_rows.append(cells)

    # Коэффициенты масштабирования к Базе по каждому виду занятий.
    def values(col: int | None) -> list[float | None]:
        if col is None:
            return [None] * len(topic_rows)
        return [_to_num(cells[col]) if col < len(cells) else None for cells in topic_rows]

    scaled: dict[str, list[str]] = {}
    totals: dict[str, float] = {}
    for key, (col, base_total) in cols.items():
        vals = values(col)
        src_sum = sum(v for v in vals if v)
        k = (base_total / src_sum) if src_sum and base_total else (1.0 if base_total else 0.0)
        out_vals: list[str] = []
        for v in vals:
            if v is None:
                out_vals.append("-")
            elif v == 0:
                out_vals.append("0")
            else:
                out_vals.append(_fmt_hours(v * k))
        if col is None and not base_total:
            out_vals = ["-"] * len(topic_rows)
        scaled[key] = out_vals
        totals[key] = float(base_total)
    order = ("lectures", "practical", "lab", "project")

    # Строки-болванки шаблона (после строки с номерами колонок «1…7») —
    # первая клонируется как образец форматирования.
    digit_idx = next(
        (i for i, r in enumerate(table.rows) if cell_text(r.cells[0]).strip() == "1"
         and cell_text(r.cells[-1]).strip() == "7"),
        2,
    )
    blanks = list(table.rows)[digit_idx + 1 :]
    if not blanks:
        return
    import copy as _copy

    proto = _copy.deepcopy(blanks[0]._tr)
    for row in blanks:
        row._tr.getparent().remove(row._tr)

    from docx.table import _Row

    def add_row(texts: list[str]) -> None:
        tr = _copy.deepcopy(proto)
        table._tbl.append(tr)
        row = _Row(tr, table)
        for c, text in enumerate(texts):
            if c >= len(row.cells):
                break
            cell = row.cells[c]
            cell.text = ""
            for pi, line in enumerate(text.split("\n")):
                paragraph = cell.paragraphs[0] if pi == 0 else cell.add_paragraph()
                run = paragraph.add_run(line)
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    for i, cells in enumerate(topic_rows):
        row_total = sum(_to_num(scaled[key][i].replace(',', '.')) or 0 for key in order)
        add_row([
            str(i + 1),
            cells[col_topic],
            *(scaled[key][i] for key in order),
            _fmt_hours(row_total),
        ])
    add_row([
        "",
        "ИТОГО",
        *(_fmt_hours(totals[key]) if totals[key] else "-" for key in order),
        _fmt_hours(sum(totals.values())),
    ])


_NUMERIC_CELL_RE = re.compile(r"[\d\s.,\-–—]+")


def _is_numeric_cell(text: str) -> bool:
    """Числовое содержимое ячейки («4», «1,6», «-») — выравнивается по центру."""
    text = text.strip()
    if not text or _NUMERIC_CELL_RE.fullmatch(text) is None:
        return False
    return any(ch.isdigit() for ch in text) or text in ("-", "–", "—")


def _apply_body_justify(docx) -> None:
    """Выравнивание по ширине для всего тела РПД (требование кафедры).

    Числовые ячейки таблиц — по центру (как в эталонах). Не трогаются:
    титульный лист (до «Москва, …» включительно), центрированные и выровненные
    вправо абзацы, шапки таблиц §3 («Вид учебной работы», «Темы (разделы)»
    до строки с номерами колонок).
    """
    untouched = (None, WD_ALIGN_PARAGRAPH.LEFT)
    body_started = False
    for p in docx.paragraphs:
        if not body_started:
            if normalize_text(p.text).startswith("москва"):
                body_started = True
            continue
        if p.text.strip() and p.paragraph_format.alignment in untouched:
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for table in docx.tables:
        if not table.rows:
            continue
        header = normalize_text(" ".join(c.text for c in table.rows[0].cells))
        skip = 0
        if "вид учебной работы" in header:
            skip = 1
        elif "темы (разделы)" in header:
            skip = next(
                (i + 1 for i, r in enumerate(table.rows)
                 if r.cells[0].text.strip() == "1" and r.cells[-1].text.strip() == "7"),
                3,
            )
        for row in table.rows[skip:]:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text.strip() and p.paragraph_format.alignment in untouched:
                        p.paragraph_format.alignment = (
                            WD_ALIGN_PARAGRAPH.CENTER
                            if _is_numeric_cell(p.text)
                            else WD_ALIGN_PARAGRAPH.JUSTIFY
                        )


def _dot(text: str) -> str:
    """Завершающая точка, как в эталонных абзацах §2."""
    return text if text.endswith(".") else text + "."


def _decapitalize(text: str) -> str:
    """Опускает первую букву («Математических…» → «математических…»);
    аббревиатуры (вторая буква заглавная) не трогает."""
    if len(text) >= 2 and text[0].isupper() and text[1].islower():
        return text[0].lower() + text[1:]
    return text


def _num(value: float) -> str | int:
    """3.0 → 3, 2.5 → 2.5 (убираем хвост .0 для целых з.е.)."""
    return int(value) if float(value).is_integer() else value
