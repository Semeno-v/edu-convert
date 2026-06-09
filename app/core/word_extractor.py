"""Чтение старых документов Word (python-docx).

Три задачи (ТЗ §4 Этап 2.1/2.3/2.5):

1. :meth:`WordExtractor.extract_index` — определить индекс дисциплины
   (титул/первая таблица).
2. :meth:`WordExtractor.extract_old_numbers` — вытащить **старые** числа из
   первой таблицы (по ключевым словам шапки) — только для отчёта о расхождениях.
3. :meth:`WordExtractor.extract` — конечный автомат по абзацам: при встрече
   заголовка включается «режим записи» блока до следующего заголовка; контент
   копируется в сериализуемый IR (:class:`ContentBlocks`).

Заголовки распознаются **по тексту** (нумерация + заглавная буква), а не по
стилям — стили в исходных документах непоследовательны (ТЗ §4 Этап 2.5).
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from app.core.exceptions import DocumentParseError
from app.core.models import (
    ContentBlock,
    ContentBlocks,
    ContentElement,
    DocType,
    ElementKind,
    OldNumbers,
    RichParagraph,
    RichRun,
    RichTable,
    RichTableCell,
    RichTableRow,
)
from app.core.normalizer import (
    as_int,
    index_from_text,
    normalize_text,
    parse_hours_ze,
)

# --------------------------------------------------------------------------- #
#  Правила сопоставления заголовков с каноническими ключами блоков.
#  ⚙️ Это главная предметная настройка: порядок важен (первое совпадение).
#  Можно расширять под особенности конкретных РПД/ФОС, не трогая логику обхода.
# --------------------------------------------------------------------------- #
BLOCK_RULES: list[tuple[str, tuple[str, ...]]] = [
    # --- РПД ---
    ("goals", ("цели освоения", "цель дисциплины")),
    ("competencies", ("компетенц", "планируемые результаты")),
    ("thematic_plan", ("тематическ", "траектори", "структура и содержание", "содержание дисциплины")),
    ("literature_main", ("основная литература",)),
    ("literature_extra", ("дополнительная литература",)),
    ("periodicals", ("периодическ",)),
    ("methodical", ("методические указания",)),
    ("internet_resources", ("интернет", "электронные образовательные", "профессиональные базы", "справочные систем")),
    ("software", ("программное обеспечение", "лицензионное")),
    ("facilities", ("материально-техническ", "технические средства", "помещения")),
    ("assessment", ("оценочные средства", "оценка качества", "система оцениван", "шкала оцениван")),
    # --- ФОС ---
    ("self_study_questions", ("самостоятельной подготовки",)),
    ("current_control", ("текущего контроля", "текущий контроль")),
    ("interim_attestation", ("промежуточной аттестации", "промежуточн")),
    ("gia", ("государственной итоговой", "итоговой аттестации", "ги(и)а")),
]

# Пробел после номера необязателен: распознаём и «3. Тема», и «3.Тема».
_HEADING_RE = re.compile(r"^\s*\d+(?:\.\d+)*\.?\s*\S")
_MAX_HEADING_LEN = 200


def _heading_key(text: str) -> str | None:
    """Возвращает канонический ключ блока по тексту заголовка (или None)."""
    norm = normalize_text(text)
    for key, keywords in BLOCK_RULES:
        if any(normalize_text(kw) in norm for kw in keywords):
            return key
    return None


def is_heading(text: str) -> bool:
    """Эвристика: нумерованный, относительно короткий заголовок раздела."""
    stripped = text.strip()
    if not stripped or len(stripped) > _MAX_HEADING_LEN:
        return False
    return bool(_HEADING_RE.match(stripped))


# --------------------------------------------------------------------------- #
#  Преобразование python-docx → IR
# --------------------------------------------------------------------------- #


def _iter_block_items(parent: DocxDocument | _Cell):
    """Итерирует абзацы и таблицы в порядке следования в документе/ячейке."""
    if isinstance(parent, DocxDocument):
        parent_elm = parent.element.body
    else:
        parent_elm = parent._tc
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def _list_level(paragraph: Paragraph) -> int | None:
    """Уровень списка из нумерации OOXML (или None)."""
    pPr = paragraph._p.pPr
    if pPr is None:
        return None
    numPr = pPr.find(qn("w:numPr"))
    if numPr is None:
        return None
    ilvl = numPr.find(qn("w:ilvl"))
    if ilvl is not None:
        val = ilvl.get(qn("w:val"))
        return as_int(val) if val is not None else 0
    return 0


# Контейнеры, внутри которых лежат раны, пропускаемые paragraph.runs:
# гиперссылки, smartTag (могут быть вложенными), простые поля.
_RUN_CONTAINER_TAGS = ("w:hyperlink", "w:smartTag", "w:fldSimple", "w:ins")


def _iter_runs(paragraph: Paragraph):
    """Раны абзаца в порядке документа, включая вложенные в w:hyperlink/w:smartTag.

    ``paragraph.runs`` python-docx отдаёт только прямые ``w:r`` — текст ссылок
    («URL: …») и smartTag-значений («2 см») иначе теряется при переносе.
    """
    container_tags = {qn(tag) for tag in _RUN_CONTAINER_TAGS}

    def walk(element):
        for child in element:
            if child.tag == qn("w:r"):
                yield Run(child, paragraph)
            elif child.tag in container_tags:
                yield from walk(child)

    yield from walk(paragraph._p)


def para_to_rich(paragraph: Paragraph) -> RichParagraph:
    """Преобразует абзац python-docx в :class:`RichParagraph` (с форматированием)."""
    runs = [
        RichRun(
            text=run.text,
            bold=bool(run.bold),
            italic=bool(run.italic),
            underline=bool(run.underline),
        )
        for run in _iter_runs(paragraph)
        if run.text
    ]
    if not runs and paragraph.text.strip():
        runs = [RichRun(text=paragraph.text)]
    style = paragraph.style.name if paragraph.style is not None else None
    return RichParagraph(runs=runs, style=style, list_level=_list_level(paragraph))


def table_to_rich(table: Table) -> RichTable:
    """Преобразует таблицу python-docx в :class:`RichTable`."""
    rows: list[RichTableRow] = []
    for row in table.rows:
        cells = [
            RichTableCell(
                paragraphs=[
                    para_to_rich(item)
                    for item in _iter_block_items(cell)
                    if isinstance(item, Paragraph)
                ]
            )
            for cell in row.cells
        ]
        rows.append(RichTableRow(cells=cells))
    return RichTable(rows=rows)


# --------------------------------------------------------------------------- #
#  Экстрактор
# --------------------------------------------------------------------------- #


class WordExtractor:
    """Чтение индекса, старых чисел и текстовых блоков из старого документа."""

    def _open(self, doc_path: Path) -> DocxDocument:
        try:
            return Document(str(doc_path))
        except Exception as exc:  # noqa: BLE001
            raise DocumentParseError(f"Не удалось открыть '{doc_path.name}': {exc}") from exc

    # ------------------------------------------------------------------ #
    #  1. Индекс
    # ------------------------------------------------------------------ #
    def extract_index(self, doc_path: Path) -> str | None:
        doc = self._open(doc_path)
        for paragraph in doc.paragraphs[:25]:
            found = index_from_text(paragraph.text)
            if found:
                return found
        for table in doc.tables[:2]:
            for row in table.rows:
                for cell in row.cells:
                    found = index_from_text(cell.text)
                    if found:
                        return found
        return None

    # ------------------------------------------------------------------ #
    #  2. Старые числа (для diff)
    # ------------------------------------------------------------------ #
    def extract_old_numbers(self, doc_path: Path) -> OldNumbers:
        doc = self._open(doc_path)
        table = self._find_hours_table(doc)
        if table is None:
            return OldNumbers()
        return self._parse_hours_table(table)

    def _find_hours_table(self, doc: DocxDocument) -> Table | None:
        for table in doc.tables:
            header_text = " ".join(
                cell.text for row in table.rows[: min(3, len(table.rows))] for cell in row.cells
            )
            n = normalize_text(header_text)
            if "трудоемкость" in n and ("лекцион" in n or "практическ" in n):
                return table
        return None

    def _parse_hours_table(self, table: Table) -> OldNumbers:
        grid = [[cell.text for cell in row.cells] for row in table.rows]
        if not grid:
            return OldNumbers()
        ncols = max(len(r) for r in grid)

        def col_text(c: int) -> str:
            return " ".join(grid[r][c] for r in range(len(grid)) if c < len(grid[r]))

        col_labels = [normalize_text(col_text(c)) for c in range(ncols)]

        def find_col(*keywords: str) -> int | None:
            for c, label in enumerate(col_labels):
                if all(kw in label for kw in keywords):
                    return c
            for c, label in enumerate(col_labels):
                if any(kw in label for kw in keywords):
                    return c
            return None

        ctrud = find_col("трудоемкость")
        if ctrud is None:
            return OldNumbers()

        # строка данных — первая ниже шапки, где трудоёмкость числовая
        data_row: list[str] | None = None
        for r in range(len(grid)):
            if ctrud < len(grid[r]):
                hours, _ze = parse_hours_ze(grid[r][ctrud])
                if hours is not None and "трудоемкость" not in normalize_text(grid[r][ctrud]):
                    data_row = grid[r]
                    break
        if data_row is None:
            return OldNumbers()

        def val(c: int | None) -> str:
            return data_row[c] if c is not None and c < len(data_row) else ""

        hours_total, ze = parse_hours_ze(val(ctrud))
        semester, _course = parse_hours_ze(val(find_col("семестр")))
        return OldNumbers(
            ze=ze,
            hours_total=as_int(hours_total) if hours_total is not None else None,
            hours_lectures=_opt_int(val(find_col("лекцион"))),
            hours_practical=_opt_int(val(find_col("практическ"))),
            hours_lab=_opt_int(val(find_col("лабораторн"))),
            hours_project=_opt_int(val(find_col("проектн"))),
            semester=int(semester) if semester is not None else None,
            control_raw=val(find_col("форма")) or None,
        )

    # ------------------------------------------------------------------ #
    #  3. Текстовые блоки (конечный автомат)
    # ------------------------------------------------------------------ #
    def extract(self, doc_path: Path, doc_type: DocType = DocType.RPD) -> ContentBlocks:
        doc = self._open(doc_path)
        result = ContentBlocks()
        self._fill_title_meta(doc, result)

        current: ContentBlock | None = None

        for item in _iter_block_items(doc):
            if isinstance(item, Paragraph):
                text = item.text.strip()
                if is_heading(text):
                    key = _heading_key(text)
                    if key is not None:
                        # Новый блок начинается ТОЛЬКО на распознанном заголовке;
                        # повторный ключ — продолжаем (merge) тот же блок, чтобы не
                        # терять контент при дублирующихся/оглавлениях-заголовках.
                        if key in result.blocks:
                            current = result.blocks[key]
                        else:
                            current = ContentBlock(key=key, title=text)
                            result.blocks[key] = current
                        continue
                    # Нумерованный, но нераспознанный заголовок (например, пункт
                    # списка «1. Дайте определение…») — это контент текущего блока,
                    # а не новый раздел: не дробим списки вопросов на части.
                if current is not None and text:
                    current.elements.append(
                        ContentElement(kind=ElementKind.PARAGRAPH, paragraph=para_to_rich(item))
                    )
            elif isinstance(item, Table):
                if current is not None:
                    current.elements.append(
                        ContentElement(kind=ElementKind.TABLE, table=table_to_rich(item))
                    )
        return result

    def _fill_title_meta(self, doc: DocxDocument, result: ContentBlocks) -> None:
        paragraphs = [p.text.strip() for p in doc.paragraphs[:30]]
        for i, text in enumerate(paragraphs):
            if not text:
                continue
            if result.index is None:
                idx = index_from_text(text)
                if idx:
                    result.index = idx
                    result.title_line = text
            low = normalize_text(text)
            if "по направлению подготовки" in low and result.direction is None:
                result.direction = re.sub(
                    r"(?i)по направлению подготовки\s*", "", text
                ).strip()
            elif low.startswith("направленности") and result.profile is None:
                result.profile = re.sub(r"(?i)направленности\s*\(профиля\)\s*", "", text).strip()
            elif "форма обучения" in low and result.form_study is None:
                result.form_study = re.sub(r"(?i)форма обучения\s*", "", text).strip()


def _opt_int(value: str) -> int | None:
    """parse_hours_ze → внешнее число (для ячеек вида «12» или «12 (…)»)."""
    outer, _ = parse_hours_ze(value)
    return as_int(outer) if outer is not None else None
